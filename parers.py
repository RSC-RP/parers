
#### PARERSv2 as of 8-2-24 ####

# input the necessary information for points 1-7

#1. Define the path to your R1 read files
a1 = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/input_data_for_RSC/truncated_MURF2_A3_WT_MGA_S5_R1_001.fastq"
b1 = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/input_data_for_RSC/truncated_MURF2_A3_L270R_MGA_S6_R1_001.fastq"

#2. Define the path to your R2 read files
a2 = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/input_data_for_RSC/truncated_MURF2_A3_WT_MGA_S5_R2_001.fastq"
b2 = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/input_data_for_RSC/truncated_MURF2_A3_L270R_MGA_S6_R2_001.fastq"

#3. Add all your R1 and R2 variables to pairwise lists
R1_list = [a1, b1]
R2_list = [a2, b2]

#4. Make a list of your cell line names (in quotes) that corresponds to the order of your R1 and R2 lists
mutant_names = ["WT MGA", "L270R MGA"]

#5. Define your control sample using its mutant name
control_sample = "WT MGA"

#6. Define the path to your input file
input_info = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/Input_Temp_11-01-24.txt"

#7. Provide the path to the FASTA file containing the raw amplicons
maxi_genes = "/data/hps/assoc/private/stuart/user/gmorto/git/CPSTUAR/stuart_parers/AmpliconsRaw.fasta"




## process input information that will apply to all cell lines ##

# load necessary packages
import re
import csv
import pandas as pd
from Bio import SeqIO
import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from collections import Counter
from Bio.Align import PairwiseAligner
from Bio.Seq import Seq
import subprocess
import gzip
#import plotly.graph_objects as go
from Bio import AlignIO
import os
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import locale
import sys
from docx.shared import Inches
import xlsxwriter
import tempfile
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement





# process information from input file
with open(input_info, 'r') as file:
    lines = file.readlines()
    if lines:
        print("File read successfully.")
    else:
        print("File is empty or could not be read.")

gene = lines[2].strip()
forward_primer = lines[5].strip()
reverse_primer = lines[8].strip()
outputfilename = lines[11].strip()
seq_orientation = lines[14].strip().upper()
cir_depth = int(lines[17].strip())
depth2 = int(lines[20].strip())
barcode_length = int(lines[23].strip())
y_axis_labels = lines[41].strip() #will still need to change this into a list to be converted for R
y_axis_increment = lines[44].strip()
y_ax_labs = y_axis_labels.split(",") #changes the y-axis labels to a list. The next step takes the min and max values and converts to integers
min_y_bound = y_ax_labs[0] #new 10/29
max_y_bound = y_ax_labs[-1] #new 10/29
output_directory_pre = lines[26].strip().replace("\\", "/")
path_to_bbmerge = lines[29].strip().replace("\\", "/")
path_to_muscle = lines[32].strip().replace("\\", "/")
path_to_r = lines[35].strip().replace("\\", "/")
path_to_r_scripts = lines[38].strip().replace("\\", "/")
path_to_bar = path_to_r_scripts + "/bargraphs_cmd_compatible.R"
path_to_bubble = path_to_r_scripts + "/bubble_plots_cmd_compatible.R"
path_to_difference = path_to_r_scripts + "/difference_plot_cmd_compatible.R"
path_to_edit_extent = path_to_r_scripts + "/editing_extent_plot_cmd_compatible.R"
path_to_edit_event = path_to_r_scripts + "/editing_event_comparison_cmd_compatible.R"
today = datetime.today().strftime('%m-%d-%Y')
sort_choice = "depth" # default to depth for now, can change to difference as needed
script_name = os.path.basename(sys.argv[0])
if output_directory_pre[-1] != "/":
    output_directory = output_directory_pre + "/"
else:
    output_directory = output_directory_pre

# create folder to store the data for all cell lines
def prompt_user(question):
    while True:
        response = input(question + " (y/n): ").strip().lower()
        if response in ['y', 'n']:
            return response == 'y'
        else:
            print("Invalid response. Please enter 'y' or 'n'.")
all_lines_folder = output_directory + outputfilename + "_" + today
if not os.path.exists(all_lines_folder):
    os.makedirs(all_lines_folder)
    print("Folder created for " + outputfilename + " in " + output_directory)
else:
    if prompt_user("Folder with the name " + all_lines_folder + " currently exists. If the program continues, existing data in the folder will be overwritten. Do you wish to proceed?"):
        print("Script will continue.")
    else:
        print("Script terminated.")
        exit()

# create a directory for the temp files. this will be deleted at the end of the run
temp_dir = tempfile.TemporaryDirectory(dir=output_directory)

# find the index for your control sample. Default set to first sample in the list
control_index = next((i for i, name in enumerate(mutant_names) if name.lower().strip() == control_sample.lower().strip()), 0)

# load gene sequences from the maxi circle sequence file
gene_sequences = {}
def load_genes(fasta_file):
    with open(fasta_file, "r") as file:
        sequence = ""
        gene_name = ""

        for line in file:
            if line.startswith(">"):
                ### Save the previous gene sequence, if any
                if gene_name and sequence:
                    gene_sequences[gene_name] = sequence
                ### Start processing a new gene sequence
                gene_name = line.strip()[1:]  # remove the leading ">"
                sequence = ""
            else:
                sequence += line.strip()

        ### Save the last gene sequence
        if gene_name and sequence:
            gene_sequences[gene_name] = sequence

load_genes(maxi_genes)

###  Retrieve the pe and fe gene sequences
# gene_names takes the gene name specified in the input file and compares that to the starting names of the gene sequences in the input FASTA file
gene_names = [key for key in gene_sequences.keys() if key.startswith(gene)]

### checks that possible genes (in gene_names) has a corollary in actual gene_sequences
if gene_names:
    for matching_gene in gene_names: # matching_gene stores the sequence names acquired by looping through gene_names
        gene_sequence = gene_sequences[matching_gene] # finds the sequence in gene_sequences FASTA file that has the name specified in gene_names
        print(f"Gene Name: {matching_gene}")
        print(f"Gene Sequence: {gene_sequence}")
else:
    print(f"No matching genes found for '{gene}' in the FASTA file.")

# compare primer sequences with pe and fe versions of the gene to see if they match
extracted_fe_sequence = ""
extracted_pe_sequence = ""
fe_sequence = None # value will be updated as long as it is found in the gene sequences FASTA
pe_sequence = None

# this code is overly complicated but takes the portion of the full sequences that we are interested in and isnt a bottle neck speed wise
if gene_names:
    for matching_gene in gene_names:
        if matching_gene.endswith("_ed"): # looking for the edited sequence
            fe_sequence = gene_sequences[matching_gene] # stores fe sequence found in the list containing all sequences, found by name
            print(fe_sequence)
            print(fe_sequence)
            if fe_sequence.find(forward_primer, fe_sequence.find(forward_primer) + len(forward_primer)) != -1 or \
                    fe_sequence.find(reverse_primer, fe_sequence.find(reverse_primer) + len(reverse_primer)) != -1:
                print(f"Forward Primer or Reverse Primer found in multiple positions in {matching_gene}")
                # if primer is found, if returns the index of first occurrence, otherwise returns a -1
                # second part in the ( ) searches for occurrences of the primer after the first, which is why you need the length of the primer
                # if primers are not found starting at the index after the length of the first primer found, then -1 is returned
                # if the primer is not found multiple times within the gene sequence then the next chunk is executed
            elif fe_sequence is not None and forward_primer in fe_sequence and reverse_primer in fe_sequence:
                fe_start_index = fe_sequence.index(forward_primer)
                fe_end_index = fe_sequence.index(reverse_primer) + len(reverse_primer)
                extracted_fe_sequence = fe_sequence[fe_start_index:fe_end_index]
                print(extracted_fe_sequence)
            else:
                print(f"Forward Primer and Reverse Primer not found in {matching_gene}")
        elif matching_gene.endswith('_pre'):
            pe_sequence = gene_sequences[matching_gene]
            if pe_sequence.find(forward_primer, pe_sequence.find(forward_primer) + len(forward_primer)) != -1 or \
                    pe_sequence.find(reverse_primer, pe_sequence.find(reverse_primer) + len(reverse_primer)) != -1:
                print(f"Forward Primer or Reverse Primer found in multiple positions in {matching_gene}")
            elif pe_sequence is not None and forward_primer in pe_sequence and reverse_primer in pe_sequence:
                pe_start_index = pe_sequence.index(forward_primer)
                pe_end_index = pe_sequence.index(reverse_primer) + len(reverse_primer)
                extracted_pe_sequence = pe_sequence[pe_start_index:pe_end_index]
            else:
                print(f"Forward Primer and Reverse Primer not found in {matching_gene}")
        else:
            print(f"Cannot find sequence for gene: {matching_gene}")
else:
    print(f"No matching genes found for '{gene}' in the FASTA file.")

print("Pre-edited sequence for analysis: " + extracted_pe_sequence)
print("Fully edited sequence for analysis: " + extracted_fe_sequence)

# get T stripped sequence
amplicon_non_t_seq = extracted_pe_sequence.replace('T', '')  # strips the Ts from the extracted pre-edited sequence

# get primer lengths
forprimlen = len(forward_primer)
revprimlen = len(reverse_primer)
barcode_plus_primer_f = forprimlen + barcode_length
barcode_plus_primer_r = revprimlen + barcode_length

#### start of per cell line loop ####
cir_files = []
bubble_files = []
bar_files = []
bar_graph_folders = []
bubble_folders = []
outputfilename_with_list = []
total_reads_list = []
num_identical_reads = []
for r1, r2, line_name in zip(R1_list, R2_list, mutant_names):

    print("Beginning processing for " + line_name)

    # create an empty folder to store all output files
    outputfilename_with_ = line_name.replace(" ", "_")
    outputfilename_with_list.append(outputfilename_with_)
    folder_name = all_lines_folder + "/" + outputfilename_with_ + "_output-files_" + today
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
        print("Folder created for " + line_name + " in " + all_lines_folder)
    else:
        if prompt_user("Sub folder with the name " + folder_name + " currently exists. If the program continues, existing data in the folder will be overwritten. "
                        "Do you wish to proceed?"):
            print("Script will continue.")
        else:
            print("Script terminated.")
            exit()

    # Process the lines in the input file
    final_fasta = folder_name + "/merged-reads_" + outputfilename_with_ + ".fasta"
    unmerged_output = folder_name + "/unmerged-reads_" + outputfilename_with_ + ".fastq"
    merged_output = os.path.join(temp_dir.name, "merged_temp.fastq")
    nont_output_aligned = os.path.join(temp_dir.name, "nont-alignments.txt")

    # determine the sequence run identifier string
    with open(r1, "r") as fastq: # changed from gzip.open to just open for rsc on 11-25-24
        first_line = str(fastq.readline().strip())
    i1 = 0
    i2 = 0
    for i, char in enumerate(first_line):
        if char == "@":
            i1 = i
        elif char == ":":
            i2 = i
            break
    seq_run_identifier = first_line[i1:i2]

    # from strictest to loosest: xstrict, ustrict, vstrict, strict, default, loose, vloose, uloose, xloose
    command = [
        # "java",
        # "-ea",
        # "-cp",
        path_to_bbmerge + "/bbmerge.sh",
        "in1=" + r1,
        "in2=" + r2,
        "out=" + merged_output,
        "outu=" + unmerged_output
    ]

    # run the command
    try:
        subprocess.run(command, check=True)
        print("BBMerge completed successfully")
    except subprocess.CalledProcessError as e:
        print("Error running BBMerge:", e)

    with open(final_fasta, 'w') as fasta:
        with open(merged_output) as fastq:
            for line in fastq:
                if line.startswith(seq_run_identifier): # previously "@"
                    header = line.strip()[1:]
                    seq = next(fastq).strip()
                    fasta.write(f'>{header}\n{seq}\n')

    print("converted to FASTA")

    merged_reads_list2 = list(SeqIO.parse(final_fasta, "fasta"))

    def count_fastq_items(filename):
        with open(filename, 'rt') as f: # gzip.
            line_count = sum(1 for line in f)
        return line_count // 4

    # determine and print out some information about how many sequences were in the original fastq files and how many were merged
    num_total_reads = count_fastq_items(r1)
    print(str(num_total_reads) + " total reads input")
    num_merged_reads = len(merged_reads_list2)
    percent_merged = round((num_merged_reads / num_total_reads) * 100, 2)
    print(str(num_merged_reads) + " (" + str(percent_merged) + "%)" " reads were merged")
    number_no_merge = num_total_reads - num_merged_reads

    # put sequences in an orientation that matches primers
    seq_with_barcodes = []
    for record in merged_reads_list2:
        if seq_orientation == "Y" or "YES":
            oriented_dict = {"seq": record.seq.reverse_complement(), "record": record.id}
            seq_with_barcodes.append(oriented_dict)
        else:
            oriented_dict = {"seq": record.seq, "record": record.id}
            seq_with_barcodes.append(oriented_dict)

    # make sure all sequences are in the same orientation before proceeding
    seq_same_orient = []
    for grou in seq_with_barcodes:
        if grou["seq"][8:revprimlen] == reverse_primer:
            seq_rc = grou["seq"].reverse_complement().replace("\n", "")
            oriented_dict = {"seq": seq_rc, "record": grou["record"]}
            seq_same_orient.append(oriented_dict)
        else:
            seq_same_orient.append(grou)

    # filter out all sequences that do not contain the forward and reverse primer
    filtered_seqs_with_barcode = []
    for forw in seq_same_orient:
        forward_seq = forw["seq"][barcode_length:barcode_plus_primer_f]
        reverse_seq = forw["seq"][-(barcode_plus_primer_r):-barcode_length]

        if forward_seq == forward_primer and reverse_seq == reverse_primer:
            filtered_seqs_with_barcode.append(forw)

    # print out some qc information regarding the number of sequences that passed the primer match filter
    num_after_primer_bin = len(filtered_seqs_with_barcode)
    post_primer_bin_percent = (len(filtered_seqs_with_barcode) / num_total_reads) * 100
    num_no_primer_bin = num_merged_reads - num_after_primer_bin

    if filtered_seqs_with_barcode:
        print("Sequences have been filtered for primer matching. " + str(num_after_primer_bin) + " (" + str((round(post_primer_bin_percent, 2))) + "%) " + "sequences remain")

    ### here starts primer deduplication
    # trim off forward and reverse barcodes
    forward_barcodes = []
    reverse_barcodes = []
    for barc in filtered_seqs_with_barcode:
        forward_barcodes.append(barc["seq"][:barcode_length])
        reverse_barcodes.append(barc["seq"][-barcode_length:])

    # concatenate the barcodes
    bar_cats = []
    for f, r in zip(forward_barcodes, reverse_barcodes):
        bar_cats.append(f + r)

    if bar_cats:
        print("Barcodes concatenated")

    # count the occurrence of each concatenated barcode
    bar_counts = Counter(bar_cats)

    # store duplicate barcodes and their counts in a dictionary
    duplicated_sequences_dict = [(seq, count) for seq, count in bar_counts.items() if count > 1]
    duplicated_sequences_dict_sorted = sorted(duplicated_sequences_dict, key=lambda x: x[1], reverse=True) # sort by highest count

    # update the filtered_seqs_with_barcodes dictionary by adding concatenated barcodes to filtered sequences dictionary
    for cats, bats in zip(filtered_seqs_with_barcode, bar_cats):
        cats["barcs_concat"] = bats

    # store all those with duplicate barcode in new list
    new_list = []
    for barco in duplicated_sequences_dict_sorted:
        new_list.append(barco[0])

    # set up a dictionary assigning an arbitrary number to each unique concatenated duplicate barcode
    mark_mapping = {value: str(i + 1) for i, value in enumerate(new_list)}

    # add the barcode assignment number to the dictionary and add a 0 if the barcode is not present in the list of duplicate barcodes
    for d in filtered_seqs_with_barcode:
        concat_barcs = d["barcs_concat"]
        d["mark"] = mark_mapping.get(concat_barcs, "0")

    # if the barcode identifier has been observed add the seq that was just iterated over to that list. If not, create a new list for that barcode
    barc_analysis_dict = {}
    for feqs in filtered_seqs_with_barcode:
        if feqs["mark"] in barc_analysis_dict:
            barc_analysis_dict[feqs["mark"]].append(feqs["seq"])
        else:
            barc_analysis_dict[feqs["mark"]] = [feqs["seq"]]

    # order the dictionary so those with the longer list values are higher
    sorted_barc_analysis_dict = sorted(barc_analysis_dict.items(), key=lambda item: len(item[1]), reverse=True)

    # convert the dictionary to a list
    barc_analysis_list = [v[1] for v in sorted_barc_analysis_dict]

    # remove sequences with non-duplicate barcodes from the list
    sorted_barc_analysis_list = barc_analysis_list[1:]

    # check through duplicate barcodes and extract all the unique sequences present between duplicates
    # this is currently only used for getting the actually_uniques list
    unique_all = []
    for listo in sorted_barc_analysis_list:
        temp_unique = set()
        unique_single = []
        for qes in listo:
            if not qes in temp_unique:
                temp_unique.add(qes)
                unique_single.append(qes)
        unique_all.append(unique_single)

    # extract all those sequences and barcodes that have multiple middle sequences
    # this is not used for anything in the current version
    actually_uniques = []
    for its in unique_all:
        if len(its) > 1:
            actually_uniques.append(its)

    # create a function that will create pairwise alignments of all sequences in a list and select out the alignment with the best score.
    # uses the Pairwise Aligner tool from Biopython that scores alignments based on the Needleman-Wunsch algorithm
    def align_sequences(seq1, seq2):
        aligner = PairwiseAligner()
        alignments = aligner.align(seq1, seq2)
        best_alignment = alignments[0]  # takes the alignment with the highest score
        aligned_seq1 = best_alignment.query
        aligned_seq2 = best_alignment.target
        return aligned_seq1, aligned_seq2

    # calculates the percent identity between aligned sequences using a base to base comparison
    def calculate_percent_identity(aligned_seq1, aligned_seq2):
        matches = sum(1 for a, b in zip(aligned_seq1, aligned_seq2) if a == b)
        identity = round(matches / len(aligned_seq1) * 100, 2)
        return identity

    # take the sequences from the list containing sequences with duplicate barcodes and create alignments and percent identities using the functions above
    percent_identities = []
    for sequences in sorted_barc_analysis_list:
        sublist_identities = []
        for i, seq1 in enumerate(sequences):
            for j in range(i + 1, len(sequences)):
                seq2 = sequences[j]
                aligned_seq1, aligned_seq2 = align_sequences(Seq(seq1), Seq(seq2))
                identity = calculate_percent_identity(aligned_seq1, aligned_seq2)
                sublist_identities.append(identity)
        percent_identities.append(sublist_identities)

    # calculate the average percent identity for all sequences that have common barcodes
    avg_percent_identities = []
    for identity in percent_identities:
        avg_percent_identities.append(sum(identity) / len(identity))

    int_identities = [int(x) for x in avg_percent_identities]

    duplicated_sequences_dict_all = [(tup[0], tup[1], flt) for tup, flt in zip(duplicated_sequences_dict_sorted, avg_percent_identities)]

    # write out barcodes and duplicate counts to a csv file
    dedup_output = folder_name + "/" + "duplicate-barcodes_" + outputfilename_with_ + ".csv"
    with open(dedup_output, "w", newline="") as outfile:
        writer = csv.writer(outfile)
        writer.writerow(["Barcodes", "Count", "% Identity"])
        for row in duplicated_sequences_dict_all:
            writer.writerow(row)

    if duplicated_sequences_dict_all:
        print("Duplicate barcodes have been counted and sorted by frequency. See 'duplicate-barcodes' csv output file.")

    # isolate unique barcodes using key value pairs
    def unique_barcodes(stuff):
        barcs = set()
        barcs_binnedd = []
        for b in stuff:
            if not b["barcs_concat"] in barcs:
                barcs.add(b["barcs_concat"])
                barcs_binnedd.append(b)
        return barcs_binnedd

    barcodes_binned = unique_barcodes(filtered_seqs_with_barcode)

    # calculate percent of sequences retained after filtering out sequences with duplicate barcodes
    barcs_binned = len(barcodes_binned)
    barcs_binned_pct = (barcs_binned / num_after_primer_bin) * 100
    barcs_binned_pct_tot = (barcs_binned / num_total_reads) * 100
    num_deduplicated = num_after_primer_bin - barcs_binned

    # trim barcodes off of the sequences and store sequences
    for seeks in barcodes_binned:
        seeks["seq"] = seeks["seq"][barcode_length:-barcode_length]

    if barcodes_binned:
        print("Sequences with duplicate barcodes have been filtered out. " + str(barcs_binned) + " (" +
              str(round(barcs_binned_pct_tot, 2)) + "%) "
              "sequences remain after filtering for primer matches and duplicate barcodes")

    # intermed1 = os.path.join(temp_dir.name, "intermed1.fasta")
    # def PrimerBinning(barcodes_binned):
    #     with open(intermed1, 'w') as output_handle: # intermed1 holds all the sequences that passed the primer binning filter. could just store this in a list but oh well
    #         for itemm in barcodes_binned:
    #             # write the sequence header
    #             output_handle.write('>' + itemm["record"] + '\n')
    #             # write the sequence on a single line
    #             output_handle.write(str(itemm["seq"]) + '\n')
    #
    # PrimerBinning(barcodes_binned)

    ### T-stripping begins ###

    # generate list of all sequences that passed primer binning
    primer_binned_seqs_list = []
    primer_binned_names_list = []
    for dict in barcodes_binned:
        primer_binned_seqs_list.append(str(dict["seq"]))
        primer_binned_names_list.append(str(dict["record"]))

    ## making a list of all the T positions in each sequence
    t_positions = []
    for deq in primer_binned_seqs_list:
        t_positions.append([i for i, nucleotide in enumerate(deq) if nucleotide == "T"])

    ### Iterate through each seq in primer_binned_seqs_list
    filtered_seqs = []
    filtered_seq_names = []
    reject_seqs = []
    reject_seqs_names = []
    for name, deq in zip(primer_binned_names_list, primer_binned_seqs_list):
        non_t_seq = deq.replace("T", "")

        ## Compare non-T seq with ref seq and filter
        if non_t_seq == amplicon_non_t_seq:
            filtered_seqs.append(deq)
            filtered_seq_names.append(name)
        else:
            reject_seqs.append(deq)
            reject_seqs_names.append(name)

    number_t_filtered_seqs = len(list(filtered_seqs))
    number_t_reject_seqs = len(list(reject_seqs))
    sum_t_filter = number_t_filtered_seqs + number_t_reject_seqs
    percent_t_filter_binned = round((number_t_filtered_seqs/barcs_binned)*100, 2)
    merged_reads_total = len(merged_reads_list2)
    percent_t_filter_total = round((number_t_filtered_seqs/num_total_reads)*100, 2)
    if sum_t_filter == barcs_binned:
        print(str(number_t_filtered_seqs) + "(" + str(percent_t_filter_total) + "%) " + "sequences remain after primer binning, deduplication, and T-stripping")
        total_reads_list.append(number_t_filtered_seqs)
    num_no_t_strip = barcs_binned - number_t_filtered_seqs

    # ## making a list of all the T positions in each T strip filtered sequence
    # t_positions_filtered = [] # stores the index of each t occurrence in each sequence that made it through T strip filtering
    # for deq in filtered_seqs:
    #     t_positions_filtered.append([i for i, nucleotide in enumerate(deq) if nucleotide == "T"])

    ### write to output file
    intermed2 = os.path.join(temp_dir.name, "intermed2.fasta")
    with open(intermed2, "w") as handle:
        for i, deq in enumerate(filtered_seqs):
            record_id = filtered_seq_names[i]
            record = f">{record_id}\n{deq}\n"
            handle.write(record)

    # Writes out the non-T mismatch sequences
    rejectststrip = folder_name + "/nonT_mismatch_" + outputfilename_with_ + ".fasta"
    with open(rejectststrip, "w") as handle:
        for i, deq in enumerate(reject_seqs):
            record_id = reject_seqs_names[i]
            record = f">{record_id}\n{deq}\n"
            handle.write(record)

    # Put non-T reject sequences into a list
    nont_reject_seqs_list = []
    ## open the fasta file and put it into a list, in string form
    with open(rejectststrip) as reject_file: # only taking the sequences that passed the primer binning test
        for record in SeqIO.parse(reject_file, "fasta"):
            nont_reject_seqs_list.append(str(record.seq))

    # strip the Ts from the non-T reject sequences
    nont_rejects_t_stripped = []
    for nont in nont_reject_seqs_list:
        non_t_seq = nont.replace("T", "")
        nont_rejects_t_stripped.append(non_t_seq)
    non_t_len = len(nont_rejects_t_stripped)

    # count the non-T rejects
    reject_counts = Counter(nont_rejects_t_stripped)

    # sort the counted non-T rejects where higher counts are at the top of the list
    sorted_non_t_rejects = {k: v for k, v in sorted(reject_counts.items(), key=lambda item: (item[1], len(item[0])), reverse=True)}

    # calculate the percent abundance of each sequence
    non_t_percentages = []
    for seq_perc in sorted_non_t_rejects.values():
        non_t_percentages.append(round((seq_perc / non_t_len)*100, 2))
    non_t_percentages2 = non_t_percentages[:cir_depth]

    # select out the top n sequences
    top_rejects = {k: sorted_non_t_rejects[k] for k in list(sorted_non_t_rejects)[:cir_depth]}

    # convert to a list
    top_rejects_list = [[key, value] for key, value in top_rejects.items()]

    arbitrary_marker = []
    for i in range(cir_depth):
        arbitrary_marker.append(str(i))

    for mork, arb in zip(top_rejects_list, arbitrary_marker):
        mork.append(str(mork[1]) + "_" + arb)

    # add the percentages to the list
    for percy, listt in zip(non_t_percentages2, top_rejects_list):
        listt.append(percy)

    # write out the top sequences to a FASTA
    rejects_fasta = os.path.join(temp_dir.name, "nont_mismatch_temp_for_MUSCLE.fasta")
    with open(rejects_fasta, "w") as f:
        f.write(">Expected non-T sequence" + "\n")
        f.write(amplicon_non_t_seq + "\n")
        for info in top_rejects_list:
            f.write(">" + str(info[2]) + " sequences (" + str(info[3]) + "%)" + "\n")
            f.write(info[0] + "\n")

    # set syntax to use MUSCLE
    command2 = [
        path_to_muscle,
        "-align", rejects_fasta,
        "-output", nont_output_aligned,
    ]

    # run MUSCLE from the command line
    try:
        subprocess.run(command2, check=True)
        print("MUSCLE completed successfully")
    except subprocess.CalledProcessError as e:
        print("Error running MUSCLE:", e)

    # Parse the output aligned sequences
    nont_reject_alignments = AlignIO.read(nont_output_aligned, "fasta")

    for record in nont_reject_alignments:
        print(record.seq + " " + record.id)

    for record in nont_reject_alignments:
        seq_id = record.id
        aligned_seq = str(record.seq)

        # Find the sublist corresponding to the sequence ID
        for sublist in top_rejects_list:
            if str(sublist[2]) == seq_id:  # Assuming seq_id is the first element in each sublist
                # Append the aligned sequence as a new item to the sublist
                sublist.append(aligned_seq)
                break  # Exit the loop once the sublist is found

    # add the percentages to the top_rejects_list as a new item for each sublist
    for sublistt, perc in zip(top_rejects_list, non_t_percentages2):
        sublistt.append(perc)

    # separate the aligned version of amplicon non-T seq
    aligned_amplicon_non_t = []
    for amp in nont_reject_alignments:
        if amp.id == "Expected":
            aligned_amplicon_non_t.append(amp.seq)
    aligned_amplicon_non_t2 = aligned_amplicon_non_t[0]

    ### writing out nont_rejects to a word doc
    nont_alignments_final = folder_name + "/" + "nonT-mismatch-alignmemts_" + outputfilename_with_ + ".docx"
    doc2 = Document()
    doc2.add_paragraph(line_name + " Most Frequent Non-T Mismatch Alignments \n")
    doc2.add_paragraph("This document contains alignments between the expected non-T amplicon sequence and the top ten most frequent sequences that did not pass the T-strip filter. "
                       "Percentages are relative to the total number of Non-T mismatches (" + str(non_t_len) + "), "
                       "not the total number of input sequences. Nucleotides are highlighted relative to the expected amplicon non-T sequence. \n")

    # write in the amplicon non-T sequence
    for record in nont_reject_alignments:
        if record.id == "Expected":
            doc2.add_paragraph(record.id + " Amplicon non-T Sequence:")
            doc2.add_paragraph(record.seq)

    # Write the alignment information to the document
    for line in top_rejects_list:
        paragraph = doc2.add_paragraph()
        paragraph_format = doc2.styles['Normal'].paragraph_format
        paragraph_format.space_after = 2
        paragraph.add_run(str(line[1]) + " (" + str(line[3]) + "%) Non-T Mismatches: \n")
        for letter1, letter2 in zip(line[4], aligned_amplicon_non_t2):
            run = paragraph.add_run(letter1)
            if letter1 != letter2 and letter1 == "G":
                run.font.color.rgb = RGBColor(255, 4, 0) # red
            elif letter1 != letter2 and letter1 == "A":
                run.font.color.rgb = RGBColor(255, 192, 0) # yellow
            elif letter1 != letter2 and letter1 == "C":
                run.font.color.rgb = RGBColor(7, 124, 199) # blue

    for paragraph in doc2.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(6)

    doc2.save(nont_alignments_final)

    # ### Run identical read condense, just do it on non-T sequence filtered dataset for this script
    # with open(intermed2, "r") as fasta_file:
    #     count_dict = {}
    #     ### read the file line by line
    #     for line in fasta_file:
    #         ### check if the line starts with a ">" character (i.e., a new sequence header)
    #         if line.startswith(">"):
    #             ### if yes, skip to the next line
    #             continue
    #         ### otherwise, remove any whitespace from the line and update the count dictionary
    #         sequence = line.strip()
    #         count_dict[sequence] = count_dict.get(sequence, 0) + 1 #this line creates a counter that tracks how many times each unique sequence is present in the FASTA file
    #         #values from this counter are stored in the count_dict list

    # condense the filtered seqs down to identical reads with frequency counts
    count_dict = {}
    for seq in filtered_seqs:
        count_dict[seq] = count_dict.get(seq, 0) + 1

    ### sort the dictionary first by frequency then defer to length of sequence if necessary
    sorted_dict = {k: v for k, v in sorted(count_dict.items(), key=lambda item: (item[1], len(item[0])), reverse=True)}
    count_dict_sum = sum(sorted_dict.values()) # get total number of sequences

    reads_percent = []
    for read in sorted_dict:
        reads_percent.append((sorted_dict[read] / count_dict_sum) * 100)
    items = sorted_dict.keys()
    item_counts = sorted_dict.values()

    # get length of each observed read
    idr_length = []
    for leg in items:
        idr_length.append(len(leg))

    # sample names here
    item_names = []
    cir_depth_range = list(range(0, len(items)))
    for i in cir_depth_range:
        item_names.append(line_name + ": Sample #" + str(i + 1))

    condense_list = []
    for s in zip(item_names, items, item_counts, idr_length, reads_percent):
        condense_list.append(s)

    cir_len = len(condense_list)
    num_identical_reads.append(cir_len)

    # pull out all identical reads that appear more than once
    idr_above_one = []
    for l in condense_list:
        if l[2] > 1:
            idr_above_one.append(l)

    # pull out all reads that occur only once
    idr_equal_one = []
    for l in condense_list:
        if l[2] == 1:
            idr_equal_one.append(l)

    # generate the condensed identical reads CSV output file
    intermed3 = folder_name + "/Condensed-Identical-Reads_" + outputfilename_with_ + ".csv"
    columns = zip(item_names, items, item_counts, idr_length, reads_percent)
    with open(intermed3, "w", newline='') as outfile:
        writer = csv.writer(outfile)
        writer.writerow(["ID", "Sequence", "Count", "Length", "Frequency"])
        writer.writerows(columns)
    cir_files.append(intermed3)

    # generate a condensed identical reads file for just those sequences that appear more than once
    idr_multiple = folder_name + "/condensed-identical-reads-above-one_" + outputfilename_with_ + "_" + ".csv"
    with open(idr_multiple, "w", newline='') as outfile2:
        writer = csv.writer(outfile2)
        writer.writerow(["ID", "Sequence", "Count", "Length", "Frequency"])
        writer.writerows(idr_above_one)

    # generate a condensed identical reads file for just those sequences that appear only once
    idr_once = folder_name + "/condensed-identical-reads-singles_" + outputfilename_with_ + "_" + ".csv"
    with open(idr_once, "w", newline='') as outfile3:
        writer = csv.writer(outfile3)
        writer.writerow(["ID", "Sequence", "Count", "Length", "Frequency"])
        writer.writerows(idr_equal_one)

    print("See condensed identical reads csv output file containing counts and percentages of all sequences")


    #############   Here begins the functions used for the right-side path giving bubble plots and bar graphs   #############

    intermed5 = os.path.join(temp_dir.name, "intermed5.txt") # will contain the number of Ts found after each A/C/G nucleotide, categorized by A/C/G
    def Tcountseq(intermed2): # uses fasta file containing sequences that passed all filters
        ## specify name of output file
        print("Generating intermediate file describing U content relative to non-U nucleotides")
        with open(intermed5, "w") as output_file:
            for record in SeqIO.parse(intermed2, "fasta"):
                seq = str(record.seq)
                seq_len = len(seq)

                # Create a list to store the counts of asterisks after each non-T nucleotide
                output_file.write(f"Sequence ID: {record.id}\n")

                for nuc in ["A", "C", "G"]:
                    res = [i.start() for i in
                           re.finditer(nuc, seq)]  # finds indexes in string of bases matching A, C or G LW
                    nuc_asterisk_counts = []
                    for i in res:  # only runs loop for where A, C or G's are located in the string/read LW
                        count = 0
                        for j in range(i + 1, seq_len):
                            if seq[j] == "T":
                                count += 1
                            else:
                                nuc_asterisk_counts.append(count)
                                count = 0
                                break
                    output_file.write(f"{nuc}: {nuc_asterisk_counts}\n")


    # #### option 2 #### probably going to keep intermediate files to avoid potential issues with overloading ram
    # intermed5_list = []
    # for seq in filtered_seqs:
    #     seq_len = len(seq)
    #
    #     acg_list = []
    #     for nuc in ["A", "C", "G"]:
    #         res = [i.start() for i in
    #                re.finditer(nuc, seq)]  # finds indexes in string of bases matching A, C or G LW
    #         nuc_asterisk_counts = []
    #         for i in res:  # only runs loop for where A, C or G's are located in the string/read LW
    #             count = 0
    #             for j in range(i + 1, seq_len):
    #                 if seq[j] == "T":
    #                     count += 1
    #                 else:
    #                     nuc_asterisk_counts.append(count)
    #                     count = 0
    #                     break
    #         acg_list.append(nuc_asterisk_counts)
    #     intermed5_list.append(acg_list)
    # #### option 2 ####

    intermed7 = os.path.join(temp_dir.name, "intermed7.csv") #this is the bubbleplot csv output but without normalized counts, order, or annotation
    def Tcountan(intermed5):
        Extrapcounts = {}
        print("STATUS: INTERMEDIATE 6 & 7")
        with open(intermed5, "r") as f, open(intermed7, "w") as output_fileLW:
            # write header in intermed7 file
            output_fileLW.write(f"Position,Nucleotide,Number,Count\n")
            for line in f:
                # If the line starts with a nucleotide, extract the counts for that nucleotide
                if line.startswith(("A:", "C:", "G:")):
                    nucleotide = line[0]
                    nucleotide_counts = [int(x) for x in line.strip().split(": ")[1].strip("[]").split(
                        ",")]  # remove non-numeric characters
                    for i, count in enumerate(nucleotide_counts):
                        if i not in Extrapcounts:
                            Extrapcounts[i] = {}
                        if nucleotide not in Extrapcounts[i]:
                            Extrapcounts[i][nucleotide] = {}
                        Extrapcounts[i][nucleotide][count] = Extrapcounts[i][nucleotide].get(count, 0) + 1

            # Write the results to the output file
            for position in Extrapcounts:
                for nucleotide in Extrapcounts[position]:
                    for count, freq in Extrapcounts[position][nucleotide].items():
                        output_fileLW.write(f"{position},{nucleotide},{count},{freq}\n")  # new intermed7 LW


    intermed8 = os.path.join(temp_dir.name, "intermed8.csv")
    def csvpatternconv(intermed7):
        df = pd.read_csv(intermed7)
        yawg = []
        yawg[:] = amplicon_non_t_seq # copies the pe T-stripped sequence from amplicon_non_t_seq to yawg
        gix = {'A': 0, 'C': 0, 'G': 0} #dictionary of nucleotides with values starting at 0
        ordpos = []
        # iterate through the nucleotides in the sequence and add 1 to the counter for each nucleotide
        for nuc in yawg:
            gix[nuc] += 1
            # append the current count of the current nucleotide to the running list
            #example output for sequence [A,C,G,A] would be [1,1,1,2]
            ordpos.append(gix[nuc])

        ordposreal = [x - 1 for x in ordpos] # normalize starting point to 0 instead of 1

        # create new dataframe with the normalized position, nucleotide, and order number
        extraporder = pd.DataFrame({"Position": ordposreal,
                                     "Nucleotide": yawg,
                                     "Order": range(1, len(ordposreal) + 1)})

        ### merge the order DataFrame with the original DataFrame to extract the corresponding rows
        merged_df = pd.merge(df, extraporder, on=["Position", "Nucleotide"], how="inner")

        ### sort the merged DataFrame by the "Order" column
        merged_df = merged_df.sort_values(by="Order")

        ## write the extracted rows to a new CSV file, excluding the "Order" column
        merged_df.to_csv(intermed8, index=False)

    # defines a function that replaces Ts with asterisks then counts the asterisks. Should just count Ts and not replace with asterisks...
    def T_count(input_string):
        brandnew_seq = ''  # make another T-stripped sequence

        brandnew_seq = input_string.replace("T", "*")

        # create a list to store the counts of asterisks after each non-T nucleotide
        asterisk_counts = []

        # count and store the asterisks in a variable
        for i, nucleotide in enumerate(brandnew_seq):
            if nucleotide in ('A', 'C', 'G'):
                count = 0
                for j in range(i + 1, len(brandnew_seq)):
                    if brandnew_seq[j] == '*':
                        count += 1
                    else:
                        asterisk_counts.append(count)
                        break
        return asterisk_counts

    # make the output files for the readnorm function
    intermed9 = os.path.join(temp_dir.name, "intermed9.csv")
    orderedcountsout = folder_name + "/" + "bubble-plot-data_" + outputfilename_with_ + '.csv'
    bubble_files.append(orderedcountsout)

    resfe = T_count(extracted_fe_sequence)  # stores counts of Ts in fully edited
    respe = T_count(extracted_pe_sequence)  # stores count of Ts in pre-edited

    # this function will generate the data for bubble plots
    def readnorm(seqcount): #intermed2 contains all sequences that passed primer binning and T stripping filters
        print("STATUS: READNORM")

        print(f"Total Sequences before normalization: {seqcount}")
        scaling = 1000000 / seqcount

        with open(intermed8, 'r') as file, open(orderedcountsout, "w",newline="") as output_ordered_fileLW:
            output_ordered_fileLW.write("Position,Nucleotide,Number,Count,Order,NewCount,Lvl1,Lvl2,Lvl3,Lvl4,Lvl5\n")
            reader = csv.DictReader(file)
            rows = list(reader)  # Read all rows into a list
            for row in rows:
                counts = int(row["Count"])
                if counts >= 1:
                    row["NewCount"] = round(counts * scaling)
                    rowv2 = row
                    order = int(rowv2["Order"])
                    number = int(rowv2["Number"])
                    if respe[order - 1] == resfe[order - 1] and number > respe[order - 1]:
                        rowv2["Lvl1"] = "NCI"
                        rowv2["Lvl2"] = "N-inc-ins"
                        rowv2["Lvl3"] = "Insertion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Insertion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif respe[order - 1] == resfe[order - 1] and number < respe[order - 1]:
                        rowv2["Lvl1"] = "NCD"
                        rowv2["Lvl2"] = "N-inc-del"
                        rowv2["Lvl3"] = "Deletion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Deletion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif respe[order - 1] == resfe[order - 1] and number == respe[order - 1]:
                        rowv2["Lvl1"] = "NE"
                        rowv2["Lvl2"] = "No edit1"
                        rowv2["Lvl3"] = "No edit2"
                        rowv2["Lvl4"] = "PE"
                        rowv2["Lvl5"] = "No edit3"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] > respe[order - 1] and number == resfe[order - 1]:
                        rowv2["Lvl1"] = "CI"
                        rowv2["Lvl2"] = "Ins ed"
                        rowv2["Lvl3"] = "Insertion"
                        rowv2["Lvl4"] = "FE"
                        rowv2["Lvl5"] = "Insertion1"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] > respe[order - 1] and respe[order - 1] < number < resfe[order - 1]:
                        rowv2["Lvl1"] = "CII"
                        rowv2["Lvl2"] = "I-inc"
                        rowv2["Lvl3"] = "Insertion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Insertion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] > respe[order - 1] and number > resfe[order - 1]:
                        rowv2["Lvl1"] = "CEI"
                        rowv2["Lvl2"] = "N-inc-ins"
                        rowv2["Lvl3"] = "Insertion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Insertion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] > respe[order - 1] and number < respe[order - 1]:
                        rowv2["Lvl1"] = "DCI"
                        rowv2["Lvl2"] = "N-inc-del"
                        rowv2["Lvl3"] = "Deletion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Deletion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] > respe[order - 1] and number == respe[order - 1]:
                        rowv2["Lvl1"] = "CINE"
                        rowv2["Lvl2"] = "No edit1"
                        rowv2["Lvl3"] = "No edit2"
                        rowv2["Lvl4"] = "PE"
                        rowv2["Lvl5"] = "No edit3"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] < respe[order - 1] and number == resfe[order - 1]:
                        rowv2["Lvl1"] = "CD"
                        rowv2["Lvl2"] = "Del ed"
                        rowv2["Lvl3"] = "Deletion"
                        rowv2["Lvl4"] = "FE"
                        rowv2["Lvl5"] = "Deletion1"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] < respe[order - 1] and resfe[order - 1] < number < respe[order - 1]:
                        rowv2["Lvl1"] = "CID"
                        rowv2["Lvl2"] = "D-inc"
                        rowv2["Lvl3"] = "Deletion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Deletion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] < respe[order - 1] and number < resfe[order - 1]:
                        rowv2["Lvl1"] = "CED"
                        rowv2["Lvl2"] = "N-inc-del"
                        rowv2["Lvl3"] = "Deletion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Deletion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] < respe[order - 1] and number > respe[order - 1]:
                        rowv2["Lvl1"] = "ICD"
                        rowv2["Lvl2"] = "N-inc-ins"
                        rowv2["Lvl3"] = "Insertion"
                        rowv2["Lvl4"] = "NC edit"
                        rowv2["Lvl5"] = "NC Insertion"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")
                    elif resfe[order - 1] < respe[order - 1] and number == respe[order - 1]:
                        rowv2["Lvl1"] = "CDNE"
                        rowv2["Lvl2"] = "No edit1"
                        rowv2["Lvl3"] = "No edit2"
                        rowv2["Lvl4"] = "PE"
                        rowv2["Lvl5"] = "No edit3"
                        output_ordered_fileLW.write(",".join(map(str, rowv2.values())) + "\n")

    # define variables for the cmp condense functions
    possible_strings1 = ["NCI", "NCD", "NE", "CI", "CII", "CEI", "DCI", "CINE", "CD", "CID", "CED", "ICD", "CDNE"] # must match the annot values in the bubble plots output
    possible_strings2 = ["N-inc-ins", "N-inc-del", "No edit1", "Ins ed", "I-inc", "Del ed", "D-inc"]
    possible_strings3 = ["Insertion", "Deletion", "No edit2"]
    possible_strings4 = ["NC edit", "PE", "FE"]
    possible_strings5 = ["No edit3", "Insertion1", "Deletion1", "NC Insertion", "NC Deletion"]
    possible_strings = possible_strings1 + possible_strings2 + possible_strings3 + possible_strings4 + possible_strings5
    baroutput = folder_name + "/bar-graph-data_" + outputfilename_with_ + ".csv"
    bar_files.append(baroutput)

    def CMPcondense(orderedcountsout):
        condensed_rows = {}
        with open(orderedcountsout, "r") as file:
            reader = csv.DictReader(file)
            for row in reader:
                nucleotide = str(row["Nucleotide"])
                order = int(row["Order"])
                count = int(row["NewCount"])
                annot1 = str(row["Lvl1"])
                annot2 = str(row["Lvl2"])
                annot3 = str(row["Lvl3"])
                annot4 = str(row["Lvl4"])
                annot5 = str(row["Lvl5"])
                annot = [annot1, annot2, annot3, annot4, annot5]
                for anno in annot:
                    if anno in possible_strings:
                        key = order
                        if key in condensed_rows:
                            # Update the counts for the matching order
                            for string in possible_strings:
                                if anno == string:
                                    condensed_rows[key][string] += count # create a counter for each editing event
                        else:
                            # Create a new entry for the order
                            condensed_rows[key] = {string: 0 for string in possible_strings}
                            condensed_rows[key]["Nucleotide"] = nucleotide
                            condensed_rows[key]["Order"] = order
                            condensed_rows[key][anno] = count

        condensed_rows_list = list(condensed_rows.values())

        with open(baroutput, "w", newline="") as file:
            fieldnames = ["Nucleotide", "Order"] + possible_strings
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()

            # Write each condensed row to the output CSV
            for row in condensed_rows_list:
                writer.writerow(row)

        print("Bar graph data written out to " + baroutput)

    #### write out alignments for just the cell line currently being investigated ####

    # store data in a list of tuples so it is easy to loop through and extract the top n designated in the input file
    reads_percs_counts_list = []
    with open(intermed3, "r") as csvfile:
        csv_reader = csv.reader(csvfile)
        for row in csv_reader:
            reads_percs_counts_list.append(row)

    adapted_list = reads_percs_counts_list[1:cir_depth + 1]

    all_sequences3 = []
    for read in adapted_list:
        all_sequences3.append(read[1])

    all_sequences2 = [extracted_pe_sequence, extracted_fe_sequence] + all_sequences3

    es_depth = depth2 - 1

    # determine scaling value to normalize read counts to reads per million reads
    rpmr_scaling = 1000000 / number_t_filtered_seqs

    # generate sequences with Ts between non-Ts
    asterisk_counts2 = []
    for asterisk in all_sequences2:
        pe_asterisk_counts = []
        pe_count = 0
        for char in asterisk:
            if char != 'T':
                if pe_count > 0:
                    # Append the count before appending the character
                    pe_asterisk_counts.append(pe_count)  # Convert count to string
                    pe_count = 0
                pe_asterisk_counts.append(char)
            else:
                pe_count += 1
        if pe_count > 0:
            # Append the count before ending the loop
            pe_asterisk_counts.append(pe_count)  # Convert count to string
        asterisk_counts2.append(pe_asterisk_counts)

    # create new sequence containing locations of ACG nucleotides and asterisk counts
    location_tracker2 = []
    for amount in asterisk_counts2:
        pe_counter = {'A': 0, 'C': 0, 'G': 0}
        pe_acg_positions = []
        for nucleotide in amount:
            if nucleotide in pe_counter:
                pe_counter[nucleotide] += 1
                pe_acg_positions.append(f"{pe_counter[nucleotide]}{nucleotide.lower()}")
            else:
                pe_acg_positions.append(nucleotide)
        location_tracker2.append(pe_acg_positions)

    # t strip the sequence
    non_t_seq = []
    for character in extracted_pe_sequence:
        if character != "T":
            non_t_seq.append(character)
    non_t_seq_string = ''.join(non_t_seq)

    # find locations of ACG nucleotides and asterisk counts in non-T sequence
    nont_counter = {'A': 0, 'C': 0, 'G': 0}
    nont_positions = []
    for nucleotide in non_t_seq_string:
        if nucleotide in nont_counter:
            nont_counter[nucleotide] += 1
            nont_positions.append(f"{nont_counter[nucleotide]}{nucleotide.lower()}")
        else:
            nont_positions.append(nucleotide)

    # select out positions just 5' of Ts and counts of Ts at each site
    t_pos2 = []
    t_counts2 = []
    for position in location_tracker2:
        pe_five_prime_positions = []
        pe_t_counts = []
        for i, pos in enumerate(position):
            if isinstance(pos, int):
                pe_t_counts.append(str(pos))
                if i > 0:  # Check if there's an item before pos
                    pe_five_prime_positions.append(position[i - 1])
        t_pos2.append(pe_five_prime_positions)
        t_counts2.append(pe_t_counts)

    if t_counts2 and t_pos2:
        print("T counts and positions have been identified for all sequences")

    # concatenate the T locations and counts for each position
    list_concat2 = []
    for post, coun in zip(t_pos2, t_counts2):
        concatenated_list = []
        for p, c in zip(post, coun):
            concatenated_list.append(p + str(c))
        list_concat2.append(concatenated_list)

    # identify all T positions across all sequences
    all_positions_list2 = []
    for post in t_pos2:
        for identifier in post:
            if identifier not in all_positions_list2:
                all_positions_list2.append(identifier)

    # order the all_positions_list sequentially
    all_positions2 = []
    for spot in nont_positions:
        if spot in all_positions_list2:
            all_positions2.append(spot)

    # create a new list where the T counts for every position (even if 0) are noted. All sublists are now the same length
    all_pos_and_count = []
    for sublist in list_concat2:
        pos_and_count = []
        for position in all_positions2:
            found = False
            for item in sublist:
                if position == item[:len(position)]:
                    pos_and_count.append(item)
                    found = True
                    break
            if not found:
                pos_and_count.append(position + '0')
        all_pos_and_count.append(pos_and_count)

    # chop out the t counts for each position into a new list
    t_counts_new2 = []
    for sublist in all_pos_and_count:
        slices = []
        for sub in sublist:
            letter_index = None  # Initialize as None
            if "g" in sub:
                letter_index = sub.index("g")
            elif "a" in sub:
                letter_index = sub.index("a")
            elif "c" in sub:
                letter_index = sub.index("c")
            slices.append(sub[letter_index + 1:])  # add 1 to get characters after the letter
        t_counts_new2.append(slices)

    # find maximum number of Ts for each position
    transposed_data2 = list(zip(*t_counts_new2))  # consolidate the t counts for each position into a separate sublist
    result2 = [[int(x) for x in sublist] for sublist in transposed_data2]  # convert counts from strings to integers
    t_max2 = [max(step) for step in result2]  # find maximum number Ts for each position

    # store the difference in number of Ts at each site editing site
    t_difference2 = []
    for newt in t_counts_new2:
        t_dif_temp = []
        for count1, count2 in zip(newt, t_max2):
            t_dif_temp.append(int(count2) - int(count1))
        t_difference2.append(t_dif_temp)

    # write out the aligned sequence
    aligned_seqs2 = []
    tee_index2 = 0  # these indices keep track of what locations are getting Ts and - added
    dif_index2 = 0
    # adds Ts or - to each T location identifier
    for t_count, dash_count in zip(t_counts_new2, t_difference2):
        single_seq_align = []
        for acg in nont_positions:
            if acg in all_positions2:
                tee = t_count[tee_index2]
                dif = dash_count[dif_index2]
                single_seq_align.append(acg + ("-" * dif) + ("T" * int(tee)))
                tee_index2 += 1  # Move to the next index for t_count
                dif_index2 += 1  # Move to the next index for dash_count
            else:
                single_seq_align.append(acg)
        aligned_seqs2.append(single_seq_align)

        # reset the indices if they exceed the length of t_count and dash_count. makes sure the t_counts and t_differences are applying to the correct location identifier
        if tee_index2 >= len(t_count):
            tee_index2 = 0
        if dif_index2 >= len(dash_count):
            dif_index2 = 0

    # remove the location numbers while keeping letters and -
    almost_aligned_seqs2 = []
    for almost_there in aligned_seqs2:
        single_final = []
        for item in almost_there:
            # remove location numbers and change letters to uppercase
            processed_item = ''.join(char for char in item if not char.isdigit()).upper()
            single_final.append(processed_item)
        almost_aligned_seqs2.append(single_final)

    # concatenate everything together for a final sequence with dashes and Ts
    final_aligned_seqs2 = []
    for joining in almost_aligned_seqs2:
        final_aligned_seqs2.append(''.join(joining))

    # replaces Ts with Us
    final_final_seqs22 = []
    for tees in final_aligned_seqs2:
        final_final_seqs22.append(tees.replace("T", "U"))

    # trim primers off sequences (except for a couple bases in the case of U editing in primer space)
    forprimlen = len(forward_primer)
    revprimlen = len(reverse_primer)
    forprimlen2 = forprimlen - 3
    revprimlen2 = revprimlen - 3
    final_final_seqs2 = []
    for sequence in final_final_seqs22:
        final_final_seqs2.append(sequence[forprimlen2:-revprimlen2])

    # acquire length of each sequence
    pre_seq_no_prim = []
    for sequ in all_sequences2:
        pre_seq_no_prim.append(sequ[forprimlen2:-revprimlen2])
    seq_lengths = [len(x) for x in pre_seq_no_prim]

    # create unique variable for pe and fe sequences with dashes and Ts. Necessary to reference for next loop
    pe_with_dash2 = final_final_seqs2[0]
    fe_with_dash2 = final_final_seqs2[1]
    pe_with_dash22 = final_final_seqs22[0]
    fe_with_dash22 = final_final_seqs22[1]

    ok_actually_final_seqs2 = []
    for sequence in final_final_seqs2:
        one_seq = []
        for base, pase in zip(sequence, pe_with_dash2):
            if base == "U" and pase == "-":
                one_seq.append("u")
            elif pase == "U" and base == "-":
                one_seq.append("*")
            else:
                one_seq.append(base)
        ok_actually_final_seqs2.append(one_seq)

    ok_actually_final_seqs22 = []
    for sequence in final_final_seqs22:
        one_seq = []
        for base, pase in zip(sequence, pe_with_dash22):
            if base == "U" and pase == "-":
                one_seq.append("u")
            elif pase == "U" and base == "-":
                one_seq.append("*")
            else:
                one_seq.append(base)
        ok_actually_final_seqs22.append(one_seq)

    # concatenate everything together for a final sequence with -, us, Us, and *
    finally_aligned_seqs21 = []
    for joining2 in ok_actually_final_seqs2:
        finally_aligned_seqs21.append(''.join(joining2))

    finally_aligned_seqs222 = []
    for joining2 in ok_actually_final_seqs22:
        finally_aligned_seqs222.append(''.join(joining2))


    # move deletions to the 3' end of the sequence
    def move_asterisks(input_string):
        # Use a regular expression to find all segments with the pattern * followed by one or more Us
        pattern = re.compile(r'\*(U+)')
        while pattern.search(input_string):
            input_string = pattern.sub(r'\1*', input_string)
        return input_string

    finally_aligned_seqs2 = []
    for string in finally_aligned_seqs21:
        updated_string = move_asterisks(string)
        finally_aligned_seqs2.append(updated_string)

    finally_aligned_seqs22 = []
    for string in finally_aligned_seqs222:
        updated_string = move_asterisks(string)
        finally_aligned_seqs22.append(updated_string)

    if finally_aligned_seqs22:
        print("Alignments have been generated using unique sequences across all cell lines.")
    ##### end of second run #####

    ##### start of the alignment for the sequences subset by editing site depth #####
    # loop through all positions where there are Us are extract the indexes of those where there are differences (ie editing sites)
    print("Starting processing of portioned data defined by editing depth.")
    pre_ed_index = []
    for i, sublist in enumerate(
            result2):  # checks the t counts at each site to see which positions have differences ie editing sites
        if not all(x == sublist[0] for x in sublist):
            pre_ed_index.append(i)
    pre_ed_index_final = pre_ed_index[-1]  # take final index representing final editing site
    true_ed_index = pre_ed_index_final - es_depth  # find the index represented the editing depth of interest
    editing_depth = []  # extract the non-T value that correlates to your editing depth of interest
    for pos in range(true_ed_index + 1, pre_ed_index_final + 1):
        editing_depth.append(all_positions2[pos])
    ed_range = [i for i in range(true_ed_index + 1, pre_ed_index_final + 1)]

    ed_site = []  # extract number position of editing base depth
    for ind in editing_depth:
        ed_site.append(int("".join([char for char in ind if char.isdigit()])))
    ed_base = []  # extract base for editing depth
    for ind in editing_depth:
        ed_base.append("".join([char for char in ind if char.isalpha()]))

    # find the index in each unprocessed sequence where the true editing depth is
    sped_index = []
    for es, eb in zip(ed_site, ed_base):
        sped_index_temp = []
        for sequence in all_sequences2:
            base_counter = 0
            for i, base in enumerate(sequence):
                if base_counter < es:
                    if base == eb.upper():
                        base_counter += 1
                        if base_counter == es:
                            sped_index_temp.append(i)
                else:
                    break
        sped_index.append(sped_index_temp)

    # find the index in the processed sequences where the true editing depth is
    proc_index = []
    for es, eb in zip(ed_site, ed_base):
        proc_temp = 0
        base_counter = 0
        for i, base in enumerate(final_final_seqs22[0]):
            if base_counter < es:
                if base == eb.upper():
                    base_counter += 1
                    if base_counter == es:
                        proc_temp = i
            else:
                break
        proc_index.append(proc_temp)

    # extract the portion of unprocessed sequence that is designated by the editing depth
    portioned_sequences = []
    for dep in sped_index:
        portioned_sequences_temp = []
        for ed_dep, seek in zip(dep, all_sequences2):
            portion = seek[ed_dep:]
            portioned_sequences_temp.append(portion)
        portioned_sequences.append(portioned_sequences_temp)

    # extract the portion of sequence of interest from the list of processed sequences
    true_portioned_sequences = []
    for pi in proc_index:
        true_portioned_sequences_temp = []
        for sequence in finally_aligned_seqs22:
            true_portion = sequence[pi:]
            true_portioned_sequences_temp.append(true_portion)
        true_portioned_sequences.append(true_portioned_sequences_temp)

    # select out only those portions of unprocessed sequences that are unique
    portioned_sequences_unique = []
    t_counts_port_unique = []
    for dept, er in zip(portioned_sequences, ed_range):
        portioned_sequences_unique_temp = []
        t_counts_port_unique_temp = []
        for sequences, tc in zip(dept, t_counts_new2):
            if sequences not in portioned_sequences_unique_temp:
                portioned_sequences_unique_temp.append(sequences)
                t_counts_port_unique_temp.append(tc[er:])
        portioned_sequences_unique.append(portioned_sequences_unique_temp)
        t_counts_port_unique.append(t_counts_port_unique_temp)

    # select out only those portions of processed sequences that are unique
    true_portioned_sequences_unique = []
    for group in true_portioned_sequences:
        true_portioned_sequences_unique_temp = []
        for sequences in group:
            if sequences not in true_portioned_sequences_unique_temp:
                true_portioned_sequences_unique_temp.append(sequences)
        true_portioned_sequences_unique.append(true_portioned_sequences_unique_temp)

    modified_reads_percs_counts_list = reads_percs_counts_list[1:]

    # search for the portion in each sequence in each mutant
    unique_portions_list_long = []
    for subo in portioned_sequences_unique:
        unique_portions_list_long_temp = []
        for port in subo:
            temp_list2 = []
            for read_data in modified_reads_percs_counts_list:
                if port in read_data[1]:
                    temp_list2.append(read_data)
            unique_portions_list_long_temp.append(temp_list2)
        unique_portions_list_long.append(unique_portions_list_long_temp)

    # condense the information from unique portions list long
    unique_portions_list_short = []
    for subo in unique_portions_list_long:
        unique_portions_list_short_temp = []
        for sequence_list in subo:
            freq_list = []
            count_list = []
            for sub_seq in sequence_list:
                perc = float(sub_seq[4])
                coun = int(sub_seq[2])
                freq_list.append(perc)
                count_list.append(coun)
            freq_sum = round(sum(freq_list), 2)
            coun_sum = sum(count_list)
            unique_portions_list_short_temp.append([freq_sum, coun_sum])
        unique_portions_list_short.append(unique_portions_list_short_temp)

    # pair the unique portions and their respective percentages for each mutant
    unique_port_with_num = []
    for sub, subo in zip(unique_portions_list_short, true_portioned_sequences_unique):
        unique_port_with_num_temp = []
        for sec, secp in zip(subo, sub):
            unique_port_with_num_temp.append([sec, secp[0], secp[1]])
        unique_port_with_num.append(unique_port_with_num_temp)

    # reformat list so it is sorted by sequence and not mutant
    uport_list_by_seq = unique_port_with_num

    # make a list of the portioned sequences in their original order
    uport_sequences_og_order = uport_list_by_seq
    # for dep in uport_list_by_seq:
    #     uport_sequences_og_order_temp = []
    #     for seq in dep:
    #         for mut in seq:
    #             if isinstance(mut, list):
    #                 uport_sequences_og_order_temp.append(mut[0][:-revprimlen2])
    #                 break
    #     uport_sequences_og_order.append(uport_sequences_og_order_temp)

    # add normalized counts to list
    for dep in uport_list_by_seq:
        for sequence in dep:
            if isinstance(sequence, list):
                normalized_ct = int(round(int(sequence[2]) * rpmr_scaling, 0))
                sequence.append(normalized_ct)

    ## processing data for portioned sequence editing depth ##
    # extract T counts for each position based on our portion
    uport_transposed_data = []
    for listy in t_counts_port_unique:
        transposed_list = [list(x) for x in zip(*listy)]
        uport_transposed_data.append(transposed_list)
    # another transformation...
    result35 = []
    for dep in uport_transposed_data:
        transformed_list = [[int(x) for x in sublist] for sublist in dep]
        result35.append(transformed_list)
    result3 = []  # update result3 to only contain U site counts for those positions in the portion of interest
    for dep, er in zip(result35, ed_range):
        truncated_list = dep[er:]
        result3.append(truncated_list)
    pe_t_ct_port = []
    for er in ed_range:
        pe_t_ct_port.append(t_counts_new2[0][er:])

    # create a list containing the last edited site by comparing T counts to PE reference. If no editing occurs, len of amplicon is used
    last_edit_index_port = []
    for dep_main, dep_pe in zip(t_counts_port_unique, pe_t_ct_port):
        last_edit_index_port_temp = []
        for sequence in dep_main:
            id_list_temp = []
            for i, (loc, ref) in enumerate(zip(sequence, dep_pe)):
                if loc != ref:
                    id_list_temp.append(i)
            if len(id_list_temp) > 0:
                last_edit_index_port_temp.append(id_list_temp[0])
            else:
                last_edit_index_port_temp.append(len(sequence) - 1)
        last_edit_index_port.append(last_edit_index_port_temp)

    # create an arbitrary ranking system
    t_ct_at_last_edit = []
    for dep1, dep2 in zip(last_edit_index_port, t_counts_port_unique):
        t_ct_at_last_edit_temp = []
        for leip, tcpu in zip(dep1, dep2):
            t_ct_at_last_edit_temp.append(int(tcpu[leip]))
        t_ct_at_last_edit.append(t_ct_at_last_edit_temp)

    # zip the lists together
    zipped_port_lists_dep = []
    for d, de, dep in zip(last_edit_index_port, t_ct_at_last_edit, uport_list_by_seq):
        zip_list = list(zip(d, de, dep))
        zipped_port_lists_dep.append(zip_list)

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    uport_list_by_seq_sorted_pre_dep = []
    for lis in zipped_port_lists_dep:
        sorted_list = sorted(lis, key=lambda x: (-x[0], x[1]))
        uport_list_by_seq_sorted_pre_dep.append(sorted_list)

    # Extract the sorted items
    uport_list_by_seq_sorted_dep = []
    for dep in uport_list_by_seq_sorted_pre_dep:
        selection = [item[2] for item in dep]
        uport_list_by_seq_sorted_dep.append(selection)

    uport_seqs_sorted2_dep = [] # make a list with just the sequences
    for dep in uport_list_by_seq_sorted_dep:
        uport_seqs_sorted2_dep_temp = []
        for seq_list in dep:
            uport_seqs_sorted2_dep_temp.append(seq_list[0])
        uport_seqs_sorted2_dep.append(uport_seqs_sorted2_dep_temp)

    uport_seqs_sorted_dep = []
    for dep in uport_seqs_sorted2_dep:
        uport_seqs_sorted_dep_temp = []
        for sequence in dep:
            uport_seqs_sorted_dep_temp.append(sequence[:-revprimlen2])
        uport_seqs_sorted_dep.append(uport_seqs_sorted_dep_temp)

    # # determine the order of sequences in the case that you sort by depth
    # dep_order_list_port = []
    # for dep, dep2 in zip(uport_sequences_og_order, uport_seqs_sorted_dep):
    #     dep_order_list_port_temp = []
    #     for seq2 in dep:
    #         for i, sequence in enumerate(dep2):
    #             if sequence == seq2:
    #                 dep_order_list_port_temp.append(i + 1)
    #                 break
    #     dep_order_list_port.append(dep_order_list_port_temp)

    ## processing data for portioned sequence difference ##

    # extract the frequencies of the non-controls
    all_uport_seq_frequencies = []
    for dep in uport_list_by_seq_sorted_dep:
        all_uport_seq_frequencies_temp = []
        for cir in dep:
            non_seq_frequencies_temp = []
            if isinstance(cir, list):
                non_seq_frequencies_temp.append(cir[3])
            else:
                non_seq_frequencies_temp.append(0)
            all_uport_seq_frequencies_temp.append(non_seq_frequencies_temp)
        all_uport_seq_frequencies.append(all_uport_seq_frequencies_temp)

    # # find the absolute difference between the number of reads in the control sample and the rest. Also determine total difference across mutants
    # total_seq_dif_port = []
    # seq_dif_port = []
    # for dep1 in all_uport_seq_frequencies:
    #     total_seq_dif_port_temp = []
    #     seq_dif_port_temp = []
    #     for cir_ex in dep1:
    #         mut_dif = []
    #         for mut in cir_ex:
    #             mut_dif.append(abs(mut - cir_con))
    #         total_seq_dif_port_temp.append(sum(mut_dif))
    #         seq_dif_port_temp.append(mut_dif)
    #     total_seq_dif_port.append(total_seq_dif_port_temp)
    #     seq_dif_port.append(seq_dif_port_temp)

    # combined_uport_by_dif = []
    # for d1, d2 in zip(total_seq_dif_port, uport_list_by_seq):
    #     zip_list = list(zip(d1, d2))
    #     combined_uport_by_dif.append(zip_list)

    # uport_sorted_by_dif_pre = []
    # for d1 in combined_uport_by_dif:
    #     sorted_list = sorted(d1, key=lambda x: x[0], reverse=True)
    #     uport_sorted_by_dif_pre.append(sorted_list)

    # uport_list_by_seq_sorted_dif = []
    # for dep in uport_sorted_by_dif_pre:
    #     sorted_list = [x[1] for x in dep]
    #     uport_list_by_seq_sorted_dif.append(sorted_list)

    # uport_seqs_sorted2_dif = []
    # for dep in uport_list_by_seq_sorted_dif:
    #     uport_seqs_sorted2_dif_temp = []
    #     for seq_list in dep:
    #         uport_seqs_sorted2_dif_temp.append(seq_list[0][0])
    #     uport_seqs_sorted2_dif.append(uport_seqs_sorted2_dif_temp)
    #
    # uport_seqs_sorted_dif = []
    # for dep in uport_seqs_sorted2_dif:
    #     uport_seqs_sorted_dif_temp = []
    #     for sequence in dep:
    #         uport_seqs_sorted_dif_temp.append(sequence[:-revprimlen2])
    #     uport_seqs_sorted_dif.append(uport_seqs_sorted_dif_temp)
    #
    # # determine the order of sequences in the case that you sort by difference
    # dif_order_list_port = []
    # for dep1, dep2 in zip(uport_sequences_og_order, uport_seqs_sorted_dif):
    #     dif_order_list_port_temp = []
    #     for seq2 in dep1:
    #         for i, sequence in enumerate(dep2):
    #             if sequence == seq2:
    #                 dif_order_list_port_temp.append(i + 1)
    #                 break
    #     dif_order_list_port.append(dif_order_list_port_temp)
    #
    # combined_uport_by_dif_final = []
    # for a, b, c, d, e in zip(total_seq_dif_port, uport_list_by_seq, dep_order_list_port, dif_order_list_port,
    #                          seq_dif_port):
    #     zist = list(zip(a, b, c, d, e))
    #     combined_uport_by_dif_final.append(zist)
    zipped_port_lists_dep_final = []
    for a, b, c in zip(last_edit_index_port, t_ct_at_last_edit, uport_list_by_seq):
        zist = list(zip(a, b, c))
        zipped_port_lists_dep_final.append(zist)

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    uport_list_by_seq_sorted_pre = []
    for dep in zipped_port_lists_dep_final:
        sist = sorted(dep, key=lambda x: (-x[0], x[1]))
        uport_list_by_seq_sorted_pre.append(sist)

    # Extract the sorted items
    uport_list_by_seq_sorted = []
    # uport_depth_order_list_sorted = []
    # uport_dif_order_list_sorted = []
    # uport_dif_per_mut_sorted = []
    for dep in uport_list_by_seq_sorted_pre:
        uport_list_by_seq_sorted.append([item[2] for item in dep])
        # uport_depth_order_list_sorted.append([item[3] for item in dep])
        # uport_dif_order_list_sorted.append([item[4] for item in dep])
        # uport_dif_per_mut_sorted.append([item[5] for item in dep])

    uport_seqs_sorted2 = []
    for dep in uport_list_by_seq_sorted:
        uport_seqs_sorted2_temp = []
        for seq_list in dep:
            uport_seqs_sorted2_temp.append(seq_list[0])
        uport_seqs_sorted2.append(uport_seqs_sorted2_temp)

    uport_seqs_sorted = []
    for dep in uport_seqs_sorted2:
        uport_seqs_sorted_temp = []
        for sequence in dep:
            uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
        uport_seqs_sorted.append(uport_seqs_sorted_temp)

    # elif sort_choice == "difference":
    #     print("Sorting sequence portions by difference between mutant and control.")
    #
    #     uport_sorted_by_dif_pre = []
    #     for dep in combined_uport_by_dif_final:
    #         sist = sorted(dep, key=lambda x: x[0], reverse=True)
    #         uport_sorted_by_dif_pre.append(sist)
    #
    #     uport_list_by_seq_sorted = []
    #     uport_depth_order_list_sorted = []
    #     uport_dif_order_list_sorted = []
    #     uport_dif_per_mut_sorted = []
    #     for dep in uport_sorted_by_dif_pre:
    #         uport_list_by_seq_sorted.append([x[1] for x in dep])
    #         uport_depth_order_list_sorted.append([item[2] for item in dep])
    #         uport_dif_order_list_sorted.append([item[3] for item in dep])
    #         uport_dif_per_mut_sorted.append([item[4] for item in dep])
    #
    #     uport_seqs_sorted2 = []
    #     for dep in uport_list_by_seq_sorted:
    #         uport_seqs_sorted2_temp = []
    #         for seq_list in dep:
    #             uport_seqs_sorted2_temp.append(seq_list[0][0])
    #         uport_seqs_sorted2.append(uport_seqs_sorted2_temp)
    #
    #     uport_seqs_sorted = []
    #     for dep in uport_seqs_sorted2:
    #         uport_seqs_sorted_temp = []
    #         for sequence in dep:
    #             uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
    #         uport_seqs_sorted.append(uport_seqs_sorted_temp)
    #
    # else:
    #     print(
    #         "There may be a typo in your choice of sequence sorting option. The script will continue without intentional sorting of portioned sequences unless terminated.")
    #
    #     uport_list_by_seq_sorted = uport_list_by_seq
    #
    #     uport_seqs_sorted2 = []
    #     for dep in uport_list_by_seq_sorted:
    #         uport_seqs_sorted2_temp = []
    #         for seq_list in dep:
    #             uport_seqs_sorted2_temp.append(seq_list[0][0])
    #         uport_seqs_sorted2.append(uport_seqs_sorted2_temp)
    #
    #     uport_seqs_sorted = []
    #     for dep in uport_seqs_sorted2:
    #         uport_seqs_sorted_temp = []
    #         for sequence in dep:
    #             uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
    #         uport_seqs_sorted.append(uport_seqs_sorted_temp)
    #
    #     uport_depth_order_list_sorted = dep_order_list_port
    #     uport_dif_order_list_sorted = dif_order_list_port
    #     uport_dif_per_mut_sorted = seq_dif_port

    ##### end of the alignment for the sequences subset by editing site depth #####

    # take out sequences for final pe and fe alignments
    finally_pe2 = finally_aligned_seqs2[0]
    finally_fe2 = finally_aligned_seqs2[1]
    finally_pe22 = finally_aligned_seqs22[0]
    finally_fe22 = finally_aligned_seqs22[1]

    # create a list containing unique sequences, excluding pe and fe
    unique_sequences = finally_aligned_seqs2[2:]

    # add the final aligned seqs to each adapted_list sublist
    for subber, sequences in zip(adapted_list, unique_sequences):
        if isinstance(subber, list):
            subber.append(sequences)

    adapted_list_by_seq = adapted_list

    # take the processed sequences in their original order
    adapted_seqs = []
    for seq_list in adapted_list_by_seq:
        if isinstance(seq_list, list):
            adapted_seqs.append(seq_list[1])

    # add normalized read counts to each mutant sublist
    for sequence in adapted_list_by_seq:
        if isinstance(sequence, list):
            normalized_ct = int(round(int(sequence[2]) * rpmr_scaling, 0))
            sequence.append(normalized_ct)

    # take out the t counts in the pe sequence for use as a reference
    pe_t_ct = t_counts_new2[0]
    # create a list containing the last edited site by comparing T counts to PE reference. If no editing occurs, len of amplicon is used
    last_edit_index = []
    for sequence in t_counts_new2[2:]:
        id_list_temp = []
        for i, (loc, ref) in enumerate(zip(sequence, pe_t_ct)):
            if loc != ref:
                id_list_temp.append(i)
        if len(id_list_temp) > 0:
            last_edit_index.append(id_list_temp[0])
        else:
            last_edit_index.append(len(sequence)-1)

    # create an arbitrary ranking system
    t_ct_at_last_edit_all_uq = []
    for leip, tcpu in zip(last_edit_index, t_counts_new2[2:]):
        t_ct_at_last_edit_all_uq.append(int(tcpu[leip]))


    for lis, seq in zip(adapted_list_by_seq, finally_aligned_seqs21):
        lis.append(seq)

    # zip the lists together
    zipped_lists_all = list(zip(last_edit_index, t_ct_at_last_edit_all_uq, adapted_list_by_seq))

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    list_by_seq_sorted_pre = sorted(zipped_lists_all, key=lambda x: (-x[0], x[1]))

    # extract the sorted items
    list_by_seq_sorted_dep = [item[2] for item in list_by_seq_sorted_pre]
    unique_seqs_sorted_dep = []
    for seq_list in list_by_seq_sorted_dep:
        if isinstance(seq_list, list):
            unique_seqs_sorted_dep.append(seq_list[5])

    # find the final order of the sequences when sorted by depth while maintaining input sequence order
    depth_order_list = []
    for seq2 in adapted_seqs:
        for i, sequence in enumerate(unique_seqs_sorted_dep):
            if sequence == seq2:
                depth_order_list.append(i + 1)

    # # extract the normalized frequencies of each sequence for your control sample
    # control_seq_frequencies = []
    # for cir in adapted_list_by_seq:
    #     if isinstance(cir, list):
    #         control_seq_frequencies.append(int(cir[control_index][6]))
    #     else:
    #         control_seq_frequencies.append(0)

    # extract the frequencies of the non-controls
    all_seq_frequencies = []
    for cir in adapted_list_by_seq:
        if isinstance(cir, list):
            all_seq_frequencies.append(cir[5])
        else:
            all_seq_frequencies.append(0)

    # # find the absolute difference between the number of reads in the control sample and the rest. Also determine total difference across mutants
    # total_seq_dif = []
    # dif_per_mut = []
    # for cir_ex, cir_con in zip(all_seq_frequencies, control_seq_frequencies):
    #     mut_dif = []
    #     for mut in cir_ex:
    #         mut_dif.append(abs(mut - cir_con))
    #     total_seq_dif.append(sum(mut_dif))
    #     dif_per_mut.append(mut_dif)

    # combined_by_dif = list(zip(total_seq_dif, adapted_list_by_seq))

    # list_sorted_by_dif_pre = sorted(combined_by_dif, key=lambda x: x[0], reverse=True)
    #
    # list_by_seq_sorted_dif = [x[1] for x in list_sorted_by_dif_pre]

    # unique_seqs_sorted_dif = []
    # for seq_list in list_by_seq_sorted_dif:
    #     for mut in seq_list:
    #         if isinstance(mut, list):
    #             unique_seqs_sorted_dif.append(mut[4])
    #             break

    # # determine the order of sequences in the case that you sort by difference
    # dif_order_list = []
    # for seq2 in adapted_seqs:
    #     for i, sequence in enumerate(unique_seqs_sorted_dif):
    #         if sequence == seq2:
    #             dif_order_list.append(i + 1)
    #             break

    #####  start of sorting by editing depth #####

    # # zip the lists together and sort everything by depth
    # zipped_lists_all2 = list(
    #     zip(last_edit_index, t_ct_at_last_edit_all_uq, adapted_list_by_seq, depth_order_list))

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    # list_by_seq_sorted_pre = sorted(zipped_lists_all2, key=lambda x: (-x[0], x[1]))

    # extract the sorted items
    list_by_seq_sorted = [item[2] for item in list_by_seq_sorted_pre]
    depth_order_list_sorted = []
    countern = 0
    for lis in list_by_seq_sorted:
        countern += 1
        depth_order_list_sorted.append(countern)
    # dif_order_list_sorted = [item[4] for item in list_by_seq_sorted_pre]
    # dif_per_mut_sorted = [item[5] for item in list_by_seq_sorted_pre]

    unique_seqs_sorted = []
    for seq_list in list_by_seq_sorted:
        if isinstance(seq_list, list):
            unique_seqs_sorted.append(seq_list[5])

    # ##### end of sorting by editing depth, start of sorting by difference between mutant and control #####
    # elif sort_choice == "difference":
    #     print("Sorting whole sequences by difference between mutant and control.")
    #
    #     # zip the lists together and sort everything by difference
    #     combined_by_dif2 = list(zip(total_seq_dif, adapted_list_by_seq, depth_order_list, dif_order_list))
    #
    #     # sort the list based on total difference between sequences
    #     list_sorted_pre = sorted(combined_by_dif2, key=lambda x: x[0], reverse=True)
    #
    #     list_by_seq_sorted = [x[1] for x in list_sorted_pre]
    #     depth_order_list_sorted = [x[2] for x in list_sorted_pre]
    #     dif_order_list_sorted = [x[3] for x in list_sorted_pre]
    #     dif_per_mut_sorted = [item[4] for item in list_sorted_pre]
    #
    #     unique_seqs_sorted = []
    #     for seq_list in list_by_seq_sorted:
    #         for mut in seq_list:
    #             if isinstance(mut, list):
    #                 unique_seqs_sorted.append(mut[5])
    #                 break
    #
    # #### this commented chunk will re-arrange the list so the control is the first column, but we will need to rearrange other lists e.g. mutant names if we want to use this #####
    # # control_list = []
    # # for sequ in list_by_seq_sorted:
    # #     control_list.append(sequ[control_index])
    # #
    # # experimental_samp_list = []
    # # for sequ in list_by_seq_sorted:
    # #     temp_list = []
    # #     for i, seu in enumerate(sequ):
    # #         if i == control_index:
    # #             continue
    # #         else:
    # #             temp_list.append(seu)
    # #     experimental_samp_list.append(temp_list)
    # #
    # # list_by_seq_sorted = []
    # # for con, exp_list in zip(control_list, experimental_samp_list):
    # #     new_list = [con] + exp_list
    # #     list_by_seq_sorted.append(new_list)
    #
    # ##### end of sorting by difference between mutant and control #####
    # else:
    #     print(
    #         "There is a typo in your choice of sequence sorting option. The script will continue without intentional sorting of whole sequences unless terminated.")
    #
    #     list_by_seq_sorted = adapted_list_by_seq
    #
    #     unique_seqs_sorted = []
    #     for seq_list in list_by_seq_sorted:
    #         unique_seqs_sorted.append(seq_list[0][5])
    #
    #     depth_order_list_sorted = depth_order_list
    #     dif_order_list_sorted = dif_order_list
    #     dif_per_mut_sorted = dif_per_mut

    # determine the lengths of each sequence
    seq_lengths = []
    for seek in unique_seqs_sorted:
        no_dash = seek.replace("-", "")
        no_ast = no_dash.replace("*", "")
        leng = len(no_ast)
        seq_lengths.append(leng)

    # process the sequences so non-U stretches are pasted in as periods
    print("Simplifying sequences by subbing non-U nucleotides for periods...")
    period_pattern = re.compile(r'[ACG]+')
    period_seqs = []
    for seeq in unique_seqs_sorted:
        output_string = period_pattern.sub('.', seeq)
        period_seqs.append(output_string)
    pe_period = period_pattern.sub(".", finally_pe2)
    fe_period = period_pattern.sub(".", finally_fe2)

    # extract the total percentages of sequences pulled for each cell line
    perc_list = []
    for sublist in list_by_seq_sorted:
        tempe = []
        if isinstance(sublist, str):
            continue
        else:
            final = round(float(sublist[4]) / 100, 4)
            perc_list.append(final)
    perc_sums = sum(perc_list)

    # repeat the above for absolute number of reads
    abs_count_list = []
    for sublist in list_by_seq_sorted:
        tempe = []
        if isinstance(sublist, str):
            continue
        else:
            final = int(sublist[2])
            abs_count_list.append(final)
    abs_count_sums = sum(abs_count_list)

    # repeat the above for reads per million reads
    rpmr_list = []
    for sublist in list_by_seq_sorted:
        tempe = []
        if isinstance(sublist, str):
            continue
        else:
            final = int(sublist[6])
            rpmr_list.append(final)
    rpmr_sums = sum(rpmr_list)

    ## repeat the above functions for the portioned sequences
    port_perc_sums = []
    for dep in uport_list_by_seq_sorted:
        mut_perc_sum = []
        for sublist in dep:
            if isinstance(sublist, str):
                continue
            else:
                final = round(float(sublist[1]) / 100, 4)
                mut_perc_sum.append(final)
        sumn = sum(mut_perc_sum)
        port_perc_sums.append(sumn)

    # repeat the above for absolute number of reads
    port_abs_count_sums = []
    for dep in uport_list_by_seq_sorted:
        mut_perc_sum = []
        for sublist in dep:
            if isinstance(sublist, str):
                continue
            else:
                final = int(sublist[2])
                mut_perc_sum.append(final)
        sumn = sum(mut_perc_sum)
        port_abs_count_sums.append(sumn)

    # repeat the above for reads per million reads
    port_rpmr_sums = []
    for dep in uport_list_by_seq_sorted:
        mut_perc_sum = []
        for sublist in dep:
            if isinstance(sublist, str):
                continue
            else:
                final = int(sublist[3])
                mut_perc_sum.append(final)
        sumn = sum(mut_perc_sum)
        port_rpmr_sums.append(sumn)

    print("All necessary data has been processed. Beginning to write out alignments to Excel sheet.")

    # lists for sorting headers and data
    mutant_names1 = ["Sequence", "Match?", "Length", ""] + [line_name]
    mutant_names2 = ["Sequence", "Match?", "Length", ""] + [line_name] + [
        ""] + [line_name] + [""] + [line_name]


    # create excel workbook
    excel_output = folder_name + "/alignments_" + outputfilename_with_ + '.xlsx'
    workbook = xlsxwriter.Workbook(excel_output)
    ws1 = workbook.add_worksheet("Full Sequence")
    ws3_name = "Period Seq Alignment"
    ws3 = workbook.add_worksheet(ws3_name)
    portion_length = sorted([i for i in range(1, es_depth + 1)], key=lambda x: x, reverse=True)
    ws2_name = []
    for i, (num, dep) in enumerate(zip(editing_depth, portion_length)):
        sheet_name = "Edit site " + num.upper() + " (depth = " + str(dep) + " ES)"
        ws2_name.append(sheet_name)
    sheet_calls = [None] * len(ws2_name)  # intialize list for the portioned tabs
    sheet_calls_list = []  # store the sheet calls for the portions
    for i, name in enumerate(ws2_name):
        sheet_calls[i] = workbook.add_worksheet(name)
        sheet_calls_list.append(sheet_calls[i])
    ws4 = workbook.add_worksheet("QC Info")
    worksheets = [ws1, ws3] + sheet_calls_list
    script_name = os.path.basename(sys.argv[0])
    today = datetime.today().strftime('%m-%d-%Y')

    # define some formats
    green_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#33CC33', 'bold': True})
    purple_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#3304FC', 'bold': True})
    red_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#FF0000', 'bold': True})
    cyan_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#3304FC', 'bold': True})
    default_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#000000'})  # Default color (black)
    bold_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#000000', 'bold': True})
    header_format = workbook.add_format({'font_name': 'Arial', 'align': 'center'})

    # prepare some columnn headers for each tab containing sequences
    for ws in worksheets:
        ws.write(0, 4, "Percent Frequencies", header_format)
        ws.write(0, 6, "Absolute # of Reads", header_format)
        ws.write(0, 8, "Reads per Million Reads", header_format)

    # write out whole sequence data to the first sheet in the workbook
    for col_index, names in enumerate(mutant_names2):
        ws1.write(1, col_index, names, header_format)

    for row_index, (custom_alignment, lol, lengs) in enumerate(zip(unique_seqs_sorted, list_by_seq_sorted, seq_lengths)):

        # Prepare the rich string with formatting
        rich_string = []

        if custom_alignment == finally_pe2:
            ws1.write(row_index + 2, 1, "Pre-edited", header_format)
            ws1.write(row_index + 2, 0, custom_alignment, default_format)
        else:
            # write out the sequences with color coding to the appropriate cell
            for character, pong, fong in zip(custom_alignment, finally_pe2, finally_fe2):
                if character == '*' and pong == "U" and fong == "U":
                    rich_string.append(green_format)
                elif character == "u" and fong == "-":
                    rich_string.append(purple_format)
                elif character == "u" and pong == "-" and fong == "u":
                    rich_string.append(red_format)
                elif character == "U" and pong == "U" and fong == "*":
                    rich_string.append(cyan_format)
                elif character == "*" and pong == "U" and fong == "*":
                    rich_string.append(bold_format)
                else:
                    rich_string.append(default_format)
                rich_string.append(character)
            # Write the rich string to the cell
            ws1.write_rich_string(row_index + 2, 0, *rich_string)

        mut_count_abs = []
        mut_perc = []
        mut_count_norm = []
        if isinstance(lol, list):
            mut_count_abs.append(int(lol[2]))
            mut_perc.append(round(float(lol[4]), 2) / 100)
            mut_count_norm.append((lol[6]))
        else:
            mut_count_abs.append("n.d.")
            mut_perc.append("n.d.")
            mut_count_norm.append("n.d.")

        ws1.write(row_index + 2, 2, lengs, header_format)

        for col_index2, mp in enumerate(mut_perc):
            ws1.write(row_index + 2, 4, mp, header_format)
        ws1.write(len(list_by_seq_sorted) + 2, 4, perc_sums, header_format)

        for col_index2, mc in enumerate(mut_count_abs):
            ws1.write(row_index + 2, 6, mc, header_format)
        ws1.write(len(list_by_seq_sorted) + 2, 6,int(abs_count_sums), header_format)

        for col_index2, mcn in enumerate(mut_count_norm):
            ws1.write(row_index + 2, 8, mcn, header_format)
        ws1.write(len(list_by_seq_sorted) + 2, 8, int(rpmr_sums), header_format)

        if custom_alignment == finally_fe2:
            ws1.write(row_index + 2, 1, "Fully edited", header_format)

    # write out the portioned data
    for (wsw, one, two, five, six, seven, eight) in zip(sheet_calls_list, uport_seqs_sorted,
                                                                     uport_list_by_seq_sorted,
                                                                     port_perc_sums, port_abs_count_sums,
                                                                     port_rpmr_sums, proc_index):
        for col_index, names in enumerate(mutant_names2):
            wsw.write(1, col_index, names, header_format)

        for row_index, (custom_alignment, lol) in enumerate(zip(one, two)):

            cu_no_dash = custom_alignment.replace("-", "")
            cu_no_ast = cu_no_dash.replace("*", "")
            leg = len(cu_no_ast)
            wsw.write(row_index + 2, 2, leg, header_format)

            # Prepare the rich string with formatting
            rich_string = []

            if custom_alignment == finally_pe2[eight:-revprimlen2]:
                wsw.write(row_index + 2, 1, "Matches pre-edited start", header_format)
                wsw.write(row_index + 2, 0, custom_alignment, default_format)
            else:
                # write out the sequences with color coding to the appropriate cell
                for character, pong, fong in zip(custom_alignment, finally_pe22[eight:-revprimlen2],
                                                 finally_fe22[eight:-revprimlen2]):
                    if character == '*' and pong == "U" and fong == "U":
                        rich_string.append(green_format)
                    elif character == "u" and fong == "-":
                        rich_string.append(purple_format)
                    elif character == "u" and pong == "-" and fong == "u":
                        rich_string.append(red_format)
                    elif character == "U" and pong == "U" and fong == "*":
                        rich_string.append(cyan_format)
                    elif character == "*" and pong == "U" and fong == "*":
                        rich_string.append(bold_format)
                    else:
                        rich_string.append(default_format)
                    rich_string.append(character)
                # Write the rich string to the cell
                wsw.write_rich_string(row_index + 2, 0, *rich_string)

            mut_count = []
            mut_perc = []
            mut_count_norm_port = []
            if lol[2] > 0:
                mut_perc.append(lol[1] / 100)
                mut_count.append(int(lol[2]))
                mut_count_norm_port.append(int(round(lol[2] * rpmr_scaling, 0)))
            else:
                mut_perc.append("n.d.")
                mut_count.append("n.d.")
                mut_count_norm_port.append("n.d.")

            for col_index2, mp in enumerate(mut_perc):
                wsw.write(row_index + 2, 4, mp, header_format)
            wsw.write(len(one) + 2, 4, five, header_format)

            for col_index2, mc in enumerate(mut_count):
                wsw.write(row_index + 2, 6, mc, header_format)
            wsw.write(len(one) + 2, 6, int(six), header_format)

            for col_index2, mcn in enumerate(mut_count_norm_port):
                wsw.write(row_index + 2, 8, mcn, header_format)
            wsw.write(len(one) + 2, 8, int(seven), header_format)

            if custom_alignment == finally_fe2[eight:-revprimlen2]:
                wsw.write(row_index + 2, 1, "Matches fully edited start", header_format)

    # write out period sequence data to the third sheet in the workbook
    for col_index, names in enumerate(mutant_names2):
        ws3.write(1, col_index, names, header_format)

    for row_index, (custom_alignment, lol) in enumerate(zip(period_seqs, list_by_seq_sorted)):

        up_u_count = custom_alignment.count("U")
        low_u_count = custom_alignment.count("u")
        u_sum = up_u_count + low_u_count
        ws3.write(row_index + 2, 2, u_sum, header_format)

        # Prepare the rich string with formatting
        rich_string = []

        if custom_alignment == pe_period:
            ws3.write(row_index + 2, 1, "Pre-edited", header_format)
            ws3.write(row_index + 2, 0, custom_alignment, default_format)
        else:
            # write out the sequences with color coding to the appropriate cell
            for character, pong, fong in zip(custom_alignment, pe_period, fe_period):
                if character == '*' and pong == "U" and fong == "U":
                    rich_string.append(green_format)
                elif character == "u" and fong == "-":
                    rich_string.append(purple_format)
                elif character == "u" and pong == "-" and fong == "u":
                    rich_string.append(red_format)
                elif character == "U" and pong == "U" and fong == "*":
                    rich_string.append(cyan_format)
                elif character == "*" and pong == "U" and fong == "*":
                    rich_string.append(bold_format)
                else:
                    rich_string.append(default_format)
                rich_string.append(character)
            # Write the rich string to the cell
            ws3.write_rich_string(row_index + 2, 0, *rich_string)

        mut_count = []
        mut_perc = []
        mut_count_norm = []
        if isinstance(lol, list):
            mut_count.append(int(lol[2]))
            mut_perc.append(round(float(lol[4]), 2) / 100)
            mut_count_norm.append(round(int(lol[2]) * rpmr_scaling, 0))
        else:
            mut_count.append("n.d.")
            mut_perc.append("n.d.")
            mut_count_norm.append("n.d.")

        for col_index2, mp in enumerate(mut_perc):
            ws3.write(row_index + 2, 4, mp, header_format)
        ws3.write(len(list_by_seq_sorted) + 2, 4, perc_sums, header_format)

        for col_index2, mc in enumerate(mut_count):
            ws3.write(row_index + 2, 6, mc, header_format)
        ws3.write(len(list_by_seq_sorted) + 2, 6, int(abs_count_sums), header_format)

        for col_index2, mcn in enumerate(mut_count_norm):
            ws3.write(row_index + 2, 8, mcn, header_format)
        ws3.write(len(list_by_seq_sorted) + 2, 8, int(rpmr_sums), header_format)

        if custom_alignment == fe_period:
            ws3.write(row_index + 2, 1, "Fully edited", header_format)

    # fill out the QC tab
    col_headers = ["# Top sequences pulled", "Control Sample", "Sorting method", "Date ran",
                   "Script version"]
    qc_data = [depth2, mutant_names[control_index], sort_choice, today, script_name]
    for row_index, title in enumerate(col_headers):
        ws4.write(row_index, 0, title, header_format)

    for row_index, info in enumerate(qc_data):
        ws4.write(row_index, 1, info, header_format)


    # save the workbook
    workbook.close()

    # run a series of functions in required order
    seqcount = count_dict_sum
    Tcountseq(intermed2)
    Tcountan(intermed5)
    csvpatternconv(intermed7)
    readnorm(seqcount)
    CMPcondense(orderedcountsout)

    # create a dictionary to store the QC information
    total_seq_in = ["Total Sequences In", num_total_reads, "N/A", "N/A", "N/A"]
    seq_merged = ["Sequences Merged", num_total_reads, num_merged_reads, number_no_merge, percent_merged]
    primer_binning = ["Primer Binning", num_merged_reads, num_after_primer_bin, num_no_primer_bin, post_primer_bin_percent]
    deduplication = ["Deduplication", num_after_primer_bin, barcs_binned, num_deduplicated, barcs_binned_pct_tot]
    t_stripping = ["T-stripping", barcs_binned, number_t_filtered_seqs, num_no_t_strip, percent_t_filter_total]
    all_qc_list = [seq_merged, primer_binning, deduplication, t_stripping]


    # generate a sankey diagram with QC data

    # fig = go.Figure(data=[go.Sankey(
    #     node = dict(
    #         pad = 10,
    #         thickness = 10,
    #         line = dict(color = "black", width = 1),
    #         label = [f"Total<br>Sequences In:<br>{num_total_reads:,}",
    #         f"Merged<br>Reads:<br>{num_merged_reads:,}",
    #         f"Primer<br>Match:<br>{num_after_primer_bin:,}",
    #         f"Deduplication:<br>{barcs_binned:,}",
    #         f"T-Stripping:<br>{number_t_filtered_seqs:,}",
    #         f"Did Not<br>Merge:<br>{number_no_merge:,}",
    #         f"Primer<br>Mismatch:<br>{num_no_primer_bin:,}",
    #         f"UMI<br>Duplicates:<br>{num_deduplicated:,}",
    #         f"Non-U<br>Mismatch:<br>{num_no_t_strip:,}"],
    #         color = ["#A1E5FE", "#5770FF", "#AB78D0", "#FE3939", "#78D385", "#5770FF", "#AB78D0", "#FE3939", "#78D385"],
    #         x = [.2, .4, .6, .8, 1, .4, .6, .8, 1],
    #         y = [.5, .5, .5, .5, .6, 1, 1, 1, .9]
    #     ),
    #     link = dict(
    #       source = [0, 1, 2, 3, 0, 1, 2, 3],
    #       target = [1, 2, 3, 4, 5, 6, 7, 8],
    #       value = [num_merged_reads, num_after_primer_bin, barcs_binned, number_t_filtered_seqs,
    #                number_no_merge, num_no_primer_bin, num_deduplicated, num_no_t_strip],
    #         color = ["#A1EEFE", "#8FA0FF", "#D9A5FF", "#FF9494", "#A1EEFE", "#8FA0FF", "#D9A5FF", "#FF9494"]
    #   ))])
    #
    # fig.update_layout(title_text= line_name + " " + gene + " QC", font_size=18, font_color="black", height=800, width=1000)
    # sankey_image = folder_name + "/" + "sankey-plot_" + outputfilename_with_ + ".pdf"
    # fig.write_image(sankey_image)

    # writing summary document
    run_summary = folder_name + "/" + "run-summary_" + outputfilename_with_ + ".docx"
    doc3 = Document()
    locale.setlocale(locale.LC_ALL, '')
    doc3.add_heading("Run summary for " + line_name + " on " + today + "\n")
    doc3.add_paragraph("Script version: " + script_name)
    doc3.add_paragraph("First read file: " + r1)
    doc3.add_paragraph("Second read file: " + r2)
    doc3.add_paragraph("Target gene: " + gene)
    doc3.add_paragraph("Forward primer: " + forward_primer)
    doc3.add_paragraph("Reverse primer: " + reverse_primer)
    doc3.add_paragraph("Portion of pre-edited sequence used for analysis: " + extracted_pe_sequence)
    doc3.add_paragraph("Portion of fully edited sequence used for analysis: " + extracted_fe_sequence)

    # format table headers
    doc3.add_paragraph("Summary of quality control data:")
    qc_table = doc3.add_table(rows = 1, cols  = 5)
    qc_table.style = "Table Grid"
    qc_headers = qc_table.rows[0].cells
    for i, cell in enumerate(qc_headers):
        cell.text = ["Step", "Before Filter", "After Filter", "Sequences Removed", "% Sequences Remaining"][i]
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align
        cell.paragraphs[0].runs[0].bold = True  # Bold text

    # format the data
    for qc in all_qc_list:
        row = qc_table.add_row().cells
        row[0].text = qc[0]
        row[1].text = "{:n}".format(qc[1])  # format number with commas
        row[2].text = "{:n}".format(qc[2])
        row[3].text = "{:n}".format(qc[3])
        row[4].text = str(round(qc[4], 2)) + "%"

    #doc3.add_picture(sankey_image, width = Inches(6), height=Inches(4))

    doc3.add_paragraph("Entire pre-edited sequence: " + pe_sequence)
    doc3.add_paragraph("Entire fully edited sequence: " + fe_sequence)

    doc3.save(run_summary)

    sankeymatic_data = [
        {"label": "Total Reads", "value": num_merged_reads, "next_label": "Merged Reads"},
        {"label": "Merged Reads", "value": num_after_primer_bin, "next_label": "Primer Match"},
        {"label": "Primer Match", "value": barcs_binned, "next_label": "Deduplication"},
        {"label": "Deduplication", "value": number_t_filtered_seqs, "next_label": "Non-U Match"},
        {"label": "Total Reads", "value": number_no_merge, "next_label": "Did Not Merge"},
        {"label": "Merged Reads", "value": num_no_primer_bin, "next_label": "Primer Mismatch"},
        {"label": "Primer Match", "value": num_deduplicated, "next_label": "UMI Duplicates"},
        {"label": "Deduplication", "value": num_no_t_strip, "next_label": "Non-U Mismatch"}
    ]

    # Define the filename for the text file
    sankeymatic_filename = folder_name + "/" + "sankeymatic_" + outputfilename_with_ + ".txt"

    # Open the file in write mode
    with open(sankeymatic_filename, "w") as file:
        # Iterate over the data
        for item in sankeymatic_data:
            # Write each item in the specified format to the file
            file.write(f"{item['label']} [{str(item['value'])}] {item['next_label']}\n")

    # generate folder to store bar graph figures later
    bar_graph_folder = folder_name + "/bar_graph_images"
    if not os.path.exists(bar_graph_folder):
        os.makedirs(bar_graph_folder)
    bar_graph_folders.append(bar_graph_folder)

    # generate folder to store bar graph figures later
    bubble_graph_folder = folder_name + "/bubble_plot_images"
    if not os.path.exists(bubble_graph_folder):
        os.makedirs(bubble_graph_folder)
    bubble_folders.append(bubble_graph_folder)

    print("Loop completed for " + line_name)

#### end of per cell line loop ####

#### beginning processing data for all cell lines ####

# create a folder for inter-cell line analyses
inter_analysis_folder = all_lines_folder + "/" + gene + "_inter_cell_line_analyses_" + today
if not os.path.exists(inter_analysis_folder):
    os.makedirs(inter_analysis_folder)
    print("Folder created for inter-cell line analyses in " + all_lines_folder)
else:
    if prompt_user("Folder with the name " + inter_analysis_folder + " currently exists. If the program continues, existing data in the folder will be overwritten. Do you wish to proceed?"):
        print("Script will continue.")
    else:
        print("Script terminated.")
        exit()

difference_folder = inter_analysis_folder + "/" + gene + "_difference_plots_" + today
if not os.path.exists(difference_folder):
    os.makedirs(difference_folder)
    print("Folder created for inter-cell line analyses in " + all_lines_folder)
else:
    if prompt_user("Folder with the name " + difference_folder + " currently exists. If the program continues, existing data in the folder will be overwritten. Do you wish to proceed?"):
        print("Script will continue.")
    else:
        print("Script terminated.")
        exit()
editing_event_figs_folder = inter_analysis_folder + "/" + gene + "_editing_event_plots_and_tables_" + today
if not os.path.exists(editing_event_figs_folder):
    os.makedirs(editing_event_figs_folder)


print("Beginning to generate R plots...")

def list_to_r_vector(py_list):
    return ','.join(py_list)

outputfilename_with_list_str = list_to_r_vector(outputfilename_with_list)
bar_files_str = list_to_r_vector(bar_files)
bar_graph_folders_str = list_to_r_vector(bar_graph_folders)
bubble_files_str = list_to_r_vector(bubble_files)
bubble_folders_str = list_to_r_vector(bubble_folders)
cir_files_str = list_to_r_vector(cir_files)
mutant_names_str = list_to_r_vector(mutant_names)
y_ax_labs_str = list_to_r_vector(y_ax_labs)

# run the R script to generate bar plots
command3 = [
    path_to_r,
    path_to_bar,
    outputfilename_with_list_str.replace("_", " "),
    forward_primer,
    reverse_primer,
    bar_files_str,
    bar_graph_folders_str,
    today,
    gene
]
try:
    subprocess.run(command3, check=True)
    print("Bar plots have been generated")
except subprocess.CalledProcessError as e:
    print("Error generating bar plots:", e)


# run the R script to generate bubble plots
command4 = [
    path_to_r,
    path_to_bubble,
    outputfilename_with_list_str.replace("_", " "),
    forward_primer,
    reverse_primer,
    bubble_files_str,
    bubble_folders_str,
    today,
    gene
]
try:
    subprocess.run(command4, check=True)
    print("Bubble plots have been generated")
except subprocess.CalledProcessError as e:
    print("Error generating bubble plots:", e)

# run the R script to generate difference plots
command5 = [
    path_to_r,
    path_to_difference,
    outputfilename_with_list_str.replace("_", " "),
    forward_primer,
    reverse_primer,
    bar_files_str,
    difference_folder,
    today,
    str(control_index),
    y_ax_labs_str, #this and lower list items added 10/30
    y_axis_increment,
    min_y_bound,
    max_y_bound,
    gene
]
try:
    subprocess.run(command5, check=True)
    print("Difference plot has been generated")
except subprocess.CalledProcessError as e:
    print("Error generating difference plots:", e)

# run the R script to generate editing extent plots
command6 = [
    path_to_r,
    path_to_edit_extent,
    extracted_pe_sequence,
    extracted_fe_sequence,
    cir_files_str,
    inter_analysis_folder,
    today,
    gene
]
try:
    subprocess.run(command6, check=True)
    print("Editing extent plot has been generated")
except subprocess.CalledProcessError as e:
    print("Error generating difference plots:", e)

# run the R script to generate editing event plots
command7 = [
    path_to_r,
    path_to_edit_event,
    forward_primer,
    reverse_primer,
    bar_files_str,
    mutant_names_str,
    editing_event_figs_folder,
    today,
    gene
]
try:
    subprocess.run(command7, check=True)
    print("Editing event plot and summary table have been generated")
except subprocess.CalledProcessError as e:
    print("Error generating difference plots:", e)


# correct for case sensitivity and extra spaces in sort choice
sort_choice = sort_choice.strip().lower()

# store data in a list of tuples so it is easy to loop through and extract the top n designated in the input file
reads_percs_counts_list = []
for csv_file in cir_files:
    with open(csv_file, newline='') as csvfile:
        csv_reader = csv.reader(csvfile)
        temp_list = []
        for row in csv_reader:
            temp_list.append(row)
    reads_percs_counts_list.append(temp_list)

# remove the top row containing columns labels
modified_reads_percs_counts_list = [sublist[1:] for sublist in reads_percs_counts_list]

depth2 = depth2 - 1

# find the total number of sequence reads for each cell line
cell_line_tot_reads = []
for cline in modified_reads_percs_counts_list:
    count_list = []
    for read in cline:
        count_list.append(int(read[2]))
    cell_line_tot_reads.append(sum(count_list))

# determine scaling value to normalize read counts to reads per million reads
rpmr_scaling = []
for read_count in cell_line_tot_reads:
    rpmr_scaling.append(1000000 / read_count)

# extract top n sequences and related counts from the contextual aligned reads csv file
topsequences = []
counts = []
top_percents = []
sequence_id = []

# take the top n elements defined by the variable depth
for item in modified_reads_percs_counts_list:
    topsequences_temp = []
    counts_temp = []
    top_percents_temp = []
    sequence_id_temp = []

    for item2 in item[:cir_depth]:
        sequence_id_temp.append(item2[0])
        topsequences_temp.append(item2[1])
        counts_temp.append(item2[2])
        top_percents_temp.append(item2[3])

    sequence_id.append(sequence_id_temp)
    topsequences.append(topsequences_temp)
    counts.append(counts_temp)
    top_percents.append(top_percents_temp)

if top_percents:
    print("Data has been successfully extracted from the input csv files.")

# make a list of the unique top sequences
top_seq_set = []
for seq in topsequences:
    for seq2 in seq:
        if seq2 not in top_seq_set:
            top_seq_set.append(seq2)


# create a list where the top sequences across all mutants are matched to each individual mutant sublist
adapted_list = []
for la_lista, mn in zip(modified_reads_percs_counts_list, mutant_names):
    top_seq_sublists = []
    for top_seq in top_seq_set:
        found = False
        for sublista in la_lista:
            if sublista[1] == top_seq:
                top_seq_sublists.append(sublista)
                found = True
                break
        if not found:
            top_seq_sublists.append("n.d")
    adapted_list.append(top_seq_sublists)

# create a list describing the number of occurrences for each of the top sequences we are looking at. Used for fasta output
gene_names = []
for n,c,p in zip(sequence_id, counts, top_percents):
    gene_names_temp = []
    for n2, c2, p2 in zip(n, c, p):
        gene_names_temp.append(n2 + " (" + str(c2) + " Occurrences, or " + str(round(float(p2), 2)) + "% of Sequences)")
    gene_names.append(gene_names_temp)

names_as_lists = []
for ct, tp, sid in zip(counts, top_percents, sequence_id):
    names_as_lists_temp = []
    for ct2, tp2, sid2 in zip(ct, tp, sid):
        names_as_lists_temp.append([sid2, ct2, tp2])
    names_as_lists.append(names_as_lists_temp)

pe_as_list = ["Pre-edited sequence", "N/A", "N/A"]
fe_as_list = ["Fully edited sequence", "N/A", "N/A"]
names_as_lists = [pe_as_list] + [fe_as_list] + names_as_lists

# almost_all_seq = []
# for listt in topsequences:
#     for sequence in listt:
#         almost_all_seq.append(sequence)

almost_all_names = []
for listy in gene_names:
    for name in listy:
        almost_all_names.append(name)

# combine all sequences into one list
all_sequences2 = [extracted_pe_sequence, extracted_fe_sequence] + list(top_seq_set)
sequence_names = ["Pre-edited sequence", "Fully edited sequence"] + almost_all_names

if sequence_names:
    print("Beginning processing of sequences to form alignments.")

# generate sequences with Ts between non-Ts
asterisk_counts2 = []
for asterisk in all_sequences2:
    pe_asterisk_counts = []
    pe_count = 0
    for char in asterisk:
        if char != 'T':
            if pe_count > 0:
                # Append the count before appending the character
                pe_asterisk_counts.append(pe_count)  # Convert count to string
                pe_count = 0
            pe_asterisk_counts.append(char)
        else:
            pe_count += 1
    if pe_count > 0:
        # Append the count before ending the loop
        pe_asterisk_counts.append(pe_count)  # Convert count to string
    asterisk_counts2.append(pe_asterisk_counts)

# create new sequence containing locations of ACG nucleotides and asterisk counts
location_tracker2 = []
for amount in asterisk_counts2:
    pe_counter = {'A': 0, 'C': 0, 'G': 0}
    pe_acg_positions = []
    for nucleotide in amount:
        if nucleotide in pe_counter:
            pe_counter[nucleotide] += 1
            pe_acg_positions.append(f"{pe_counter[nucleotide]}{nucleotide.lower()}")
        else:
            pe_acg_positions.append(nucleotide)
    location_tracker2.append(pe_acg_positions)

# find locations of ACG nucleotides and asterisk counts in non-T sequence
nont_counter = {'A': 0, 'C': 0, 'G': 0}
nont_positions = []
for nucleotide in amplicon_non_t_seq:
    if nucleotide in nont_counter:
        nont_counter[nucleotide] += 1
        nont_positions.append(f"{nont_counter[nucleotide]}{nucleotide.lower()}")
    else:
        nont_positions.append(nucleotide)

# select out positions just 5' of Ts and counts of Ts at each site
t_pos2 = []
t_counts2 = []
for position in location_tracker2:
    pe_five_prime_positions = []
    pe_t_counts = []
    for i, pos in enumerate(position):
        if isinstance(pos, int):
            pe_t_counts.append(str(pos))
            if i > 0:  # Check if there's an item before pos
                pe_five_prime_positions.append(position[i - 1])
    t_pos2.append(pe_five_prime_positions)
    t_counts2.append(pe_t_counts)

if t_counts2 and t_pos2:
    print("T counts and positions have been identified for all sequences")

# concatenate the T locations and counts for each position
list_concat2 = []
for post, coun in zip(t_pos2, t_counts2):
    concatenated_list = []
    for p, c in zip(post, coun):
        concatenated_list.append(p + str(c))
    list_concat2.append(concatenated_list)

# identify all T positions across all sequences
all_positions_list2 = []
for post in t_pos2:
    for identifier in post:
        if identifier not in all_positions_list2:
            all_positions_list2.append(identifier)

# order the all_positions_list sequentially
all_positions2 = []
for spot in nont_positions:
    if spot in all_positions_list2:
        all_positions2.append(spot)

# create a new list where the T counts for every position (even if 0) are noted. All sublists are now the same length
all_pos_and_count = []
for sublist in list_concat2:
    pos_and_count = []
    for position in all_positions2:
        found = False
        for item in sublist:
            if position == item[:len(position)]:
                pos_and_count.append(item)
                found = True
                break
        if not found:
            pos_and_count.append(position + '0')
    all_pos_and_count.append(pos_and_count)

# chop out the t counts for each position into a new list
t_counts_new2 = []
for sublist in all_pos_and_count:
    slices = []
    for sub in sublist:
        letter_index = None  # Initialize as None
        if "g" in sub:
            letter_index = sub.index("g")
        elif "a" in sub:
            letter_index = sub.index("a")
        elif "c" in sub:
            letter_index = sub.index("c")
        slices.append(sub[letter_index + 1:])  # add 1 to get characters after the letter
    t_counts_new2.append(slices)

# find maximum number of Ts for each position
transposed_data2 = list(zip(*t_counts_new2)) # consolidate the t counts for each position into a separate sublist
result2 = [[int(x) for x in sublist] for sublist in transposed_data2] # convert counts from strings to integers
t_max2 = [max(step) for step in result2] # find maximum number Ts for each position

# store the difference in number of Ts at each site editing site
t_difference2 = []
for newt in t_counts_new2:
    t_dif_temp = []
    for count1, count2 in zip(newt, t_max2):
        t_dif_temp.append(int(count2) - int(count1))
    t_difference2.append(t_dif_temp)

# write out the aligned sequence
aligned_seqs2 = []
tee_index2 = 0  # these indices keep track of what locations are getting Ts and - added
dif_index2 = 0

# adds Ts or - to each T location identifier
for t_count, dash_count in zip(t_counts_new2, t_difference2):
    single_seq_align = []
    for acg in nont_positions:
        if acg in all_positions2:
            tee = t_count[tee_index2]
            dif = dash_count[dif_index2]
            single_seq_align.append(acg + ("-" * dif) + ("T" * int(tee)))
            tee_index2 += 1  # Move to the next index for t_count
            dif_index2 += 1  # Move to the next index for dash_count
        else:
            single_seq_align.append(acg)
    aligned_seqs2.append(single_seq_align)

    # reset the indices if they exceed the length of t_count and dash_count. makes sure the t_counts and t_differences are applying to the correct location identifier
    if tee_index2 >= len(t_count):
        tee_index2 = 0
    if dif_index2 >= len(dash_count):
        dif_index2 = 0

# remove the location numbers while keeping letters and -
almost_aligned_seqs2 = []
for almost_there in aligned_seqs2:
    single_final = []
    for item in almost_there:
        # remove location numbers and change letters to uppercase
        processed_item = ''.join(char for char in item if not char.isdigit()).upper()
        single_final.append(processed_item)
    almost_aligned_seqs2.append(single_final)

# concatenate everything together for a final sequence with dashes and Ts
final_aligned_seqs2 = []
for joining in almost_aligned_seqs2:
    final_aligned_seqs2.append(''.join(joining))

# replaces Ts with Us
final_final_seqs22 = []
for tees in final_aligned_seqs2:
    final_final_seqs22.append(tees.replace("T", "U"))

# trim primers off sequences (except for a couple bases in the case of U editing in primer space)
forprimlen = len(forward_primer)
revprimlen = len(reverse_primer)
forprimlen2 = forprimlen - 3
revprimlen2 = revprimlen - 3
final_final_seqs2 = []
for sequence in final_final_seqs22:
    final_final_seqs2.append(sequence[forprimlen2:-revprimlen2])

# acquire length of each sequence
pre_seq_no_prim = []
for sequ in all_sequences2:
    pre_seq_no_prim.append(sequ[forprimlen2:-revprimlen2])
seq_lengths = [len(x) for x in pre_seq_no_prim]

# create unique variable for pe and fe sequences with dashes and Ts. Necessary to reference for next loop
pe_with_dash2 = final_final_seqs2[0]
fe_with_dash2 = final_final_seqs2[1]
pe_with_dash22 = final_final_seqs22[0]
fe_with_dash22 = final_final_seqs22[1]

ok_actually_final_seqs2 = []
for sequence in final_final_seqs2:
    one_seq = []
    for base, pase in zip(sequence, pe_with_dash2):
        if base == "U" and pase == "-":
            one_seq.append("u")
        elif pase == "U" and base == "-":
            one_seq.append("*")
        else:
            one_seq.append(base)
    ok_actually_final_seqs2.append(one_seq)

ok_actually_final_seqs22 = []
for sequence in final_final_seqs22:
    one_seq = []
    for base, pase in zip(sequence, pe_with_dash22):
        if base == "U" and pase == "-":
            one_seq.append("u")
        elif pase == "U" and base == "-":
            one_seq.append("*")
        else:
            one_seq.append(base)
    ok_actually_final_seqs22.append(one_seq)

# concatenate everything together for a final sequence with -, us, Us, and *
finally_aligned_seqs21 = []
for joining2 in ok_actually_final_seqs2:
    finally_aligned_seqs21.append(''.join(joining2))

finally_aligned_seqs222 = []
for joining2 in ok_actually_final_seqs22:
    finally_aligned_seqs222.append(''.join(joining2))

# move deletions to the 3' end of the sequence
def move_asterisks(input_string):
    # Use a regular expression to find all segments with the pattern * followed by one or more Us
    pattern = re.compile(r'\*(U+)')
    while pattern.search(input_string):
        input_string = pattern.sub(r'\1*', input_string)
    return input_string

finally_aligned_seqs2 = []
for string in finally_aligned_seqs21:
    updated_string = move_asterisks(string)
    finally_aligned_seqs2.append(updated_string)

### use map?? ###
finally_aligned_seqs_mapped = map(move_asterisks, finally_aligned_seqs21)

finally_aligned_seqs22 = []
for string in finally_aligned_seqs222:
    updated_string = move_asterisks(string)
    finally_aligned_seqs22.append(updated_string)

if finally_aligned_seqs22:
    print("Alignments have been generated using unique sequences across all cell lines.")
##### end of second run #####

##### start of the alignment for the sequences subset by editing site depth #####
# loop through all positions where there are Us are extract the indexes of those where there are differences (ie editing sites)
print("Starting processing of portioned data defined by editing depth.")
pre_ed_index = []
for i, sublist in enumerate(result2): # checks the t counts at each site to see which positions have differences ie editing sites
    if not all(x == sublist[0] for x in sublist):
        pre_ed_index.append(i)
pre_ed_index_final = pre_ed_index[-1] # take final index representing final editing site
true_ed_index = pre_ed_index_final - depth2 # find the index represented the editing depth of interest
editing_depth = [] # extract the non-T value that correlates to your editing depth of interest
for pos in range(true_ed_index + 1, pre_ed_index_final + 1):
    editing_depth.append(all_positions2[pos])
ed_range = [i for i in range(true_ed_index + 1, pre_ed_index_final + 1)]


ed_site = [] # extract number position of editing base depth
for ind in editing_depth:
    ed_site.append(int("".join([char for char in ind if char.isdigit()])))
ed_base = [] # extract base for editing depth
for ind in editing_depth:
    ed_base.append("".join([char for char in ind if char.isalpha()]))

# find the index in each unprocessed sequence where the true editing depth is
sped_index = []
for es, eb in zip(ed_site, ed_base):
    sped_index_temp = []
    for sequence in all_sequences2:
        base_counter = 0
        for i, base in enumerate(sequence):
            if base_counter < es:
                if base == eb.upper():
                    base_counter += 1
                    if base_counter == es:
                        sped_index_temp.append(i)
            else:
                break
    sped_index.append(sped_index_temp)

# find the index in the processed sequences where the true editing depth is
proc_index = []
for es, eb in zip(ed_site, ed_base):
    proc_temp = 0
    base_counter = 0
    for i, base in enumerate(final_final_seqs22[0]):
        if base_counter < es:
            if base == eb.upper():
                base_counter += 1
                if base_counter == es:
                    proc_temp = i
        else:
            break
    proc_index.append(proc_temp)

# extract the portion of unprocessed sequence that is designated by the editing depth
portioned_sequences = []
for dep in sped_index:
    portioned_sequences_temp = []
    for ed_dep, seek in zip(dep, all_sequences2):
        portion = seek[ed_dep:]
        portioned_sequences_temp.append(portion)
    portioned_sequences.append(portioned_sequences_temp)

# extract the portion of sequence of interest from the list of processed sequences
true_portioned_sequences = []
for pi in proc_index:
    true_portioned_sequences_temp = []
    for sequence in finally_aligned_seqs22:
        true_portion = sequence[pi:]
        true_portioned_sequences_temp.append(true_portion)
    true_portioned_sequences.append(true_portioned_sequences_temp)

# select out only those portions of unprocessed sequences that are unique
portioned_sequences_unique = []
t_counts_port_unique = []
for dept, er in zip(portioned_sequences, ed_range):
    portioned_sequences_unique_temp = []
    t_counts_port_unique_temp = []
    for sequences, tc in zip(dept, t_counts_new2):
        if sequences not in portioned_sequences_unique_temp:
            portioned_sequences_unique_temp.append(sequences)
            t_counts_port_unique_temp.append(tc[er:])
    portioned_sequences_unique.append(portioned_sequences_unique_temp)
    t_counts_port_unique.append(t_counts_port_unique_temp)

# select out only those portions of processed sequences that are unique
true_portioned_sequences_unique = []
for group in true_portioned_sequences:
    true_portioned_sequences_unique_temp = []
    for sequences in group:
        if sequences not in true_portioned_sequences_unique_temp:
            true_portioned_sequences_unique_temp.append(sequences)
    true_portioned_sequences_unique.append(true_portioned_sequences_unique_temp)

# search for the portion in each sequence in each mutant
unique_portions_list_long = []
for subo in portioned_sequences_unique:
    unique_portions_list_long_temp = []
    for mutant_list in modified_reads_percs_counts_list:
        temp_list1 = []
        for port in subo:
            temp_list2 = []
            for read_data in mutant_list:
                if port in read_data[1]:
                    temp_list2.append(read_data)
            temp_list1.append(temp_list2)
        unique_portions_list_long_temp.append(temp_list1)
    unique_portions_list_long.append(unique_portions_list_long_temp)

# condense the information from unique portions list long
unique_portions_list_short = []
for subo in unique_portions_list_long:
    unique_portions_list_short_temp = []
    for mutant_list in subo:
        sequence_sl = []
        for sequence_list in mutant_list:
            freq_list = []
            count_list = []
            for sub_seq in sequence_list:
                perc = float(sub_seq[4])
                coun = int(sub_seq[2])
                freq_list.append(perc)
                count_list.append(coun)
            freq_sum = round(sum(freq_list), 2)
            coun_sum = sum(count_list)
            sequence_sl.append([freq_sum, coun_sum])
        unique_portions_list_short_temp.append(sequence_sl)
    unique_portions_list_short.append(unique_portions_list_short_temp)

# pair the unique portions and their respective percentages for each mutant
unique_port_with_num = []
for sub, subo in zip(unique_portions_list_short, true_portioned_sequences_unique):
    unique_port_with_num_temp = []
    for mut in sub:
        sl1 = []
        for subv, sec in zip(mut, subo):
            sl1.append([sec, subv[0], subv[1]])
        unique_port_with_num_temp.append(sl1)
    unique_port_with_num.append(unique_port_with_num_temp)

# reformat list so it is sorted by sequence and not mutant
uport_list_by_seq = []
for listy in unique_port_with_num:
    transposed_list = [list(x) for x in zip(*listy)]
    uport_list_by_seq.append(transposed_list)

# make a list of the portioned sequences in their original order
uport_sequences_og_order = []
for dep in uport_list_by_seq:
    uport_sequences_og_order_temp = []
    for seq in dep:
        for mut in seq:
            if isinstance(mut, list):
                uport_sequences_og_order_temp.append(mut[0][:-revprimlen2])
                break
    uport_sequences_og_order.append(uport_sequences_og_order_temp)

# add normalized counts to list
for dep in uport_list_by_seq:
    for sequence in dep:
        for mut, rpmr in zip(sequence, rpmr_scaling):
            if isinstance(mut, list):
                normalized_ct = int(round(int(mut[2]) * rpmr, 0))
                mut.append(normalized_ct)

## processing data for portioned sequence editing depth ##
# extract T counts for each position based on our portion
uport_transposed_data = []
for listy in t_counts_port_unique:
    transposed_list = [list(x) for x in zip(*listy)]
    uport_transposed_data.append(transposed_list)
# another transformation...
result35 = []
for dep in uport_transposed_data:
    transformed_list = [[int(x) for x in sublist] for sublist in dep]
    result35.append(transformed_list)
result3 = [] # update result3 to only contain U site counts for those positions in the portion of interest
for dep, er in zip(result35, ed_range):
    truncated_list = dep[er:]
    result3.append(truncated_list)
pe_t_ct_port = []
for er in ed_range:
    pe_t_ct_port.append(t_counts_new2[0][er:])

# create a list containing the last edited site by comparing T counts to PE reference. If no editing occurs, len of amplicon is used
last_edit_index_port = []
for dep_main, dep_pe in zip(t_counts_port_unique, pe_t_ct_port):
    last_edit_index_port_temp = []
    for sequence in dep_main:
        id_list_temp = []
        for i, (loc, ref) in enumerate(zip(sequence, dep_pe)):
            if loc != ref:
                id_list_temp.append(i)
        if len(id_list_temp) > 0:
            last_edit_index_port_temp.append(id_list_temp[0])
        else:
            last_edit_index_port_temp.append(len(sequence) - 1)
    last_edit_index_port.append(last_edit_index_port_temp)

# create an arbitrary ranking system
t_ct_at_last_edit = []
for dep1, dep2 in zip(last_edit_index_port, t_counts_port_unique):
    t_ct_at_last_edit_temp = []
    for leip, tcpu in zip(dep1, dep2):
        t_ct_at_last_edit_temp.append(int(tcpu[leip]))
    t_ct_at_last_edit.append(t_ct_at_last_edit_temp)

# zip the lists together
zipped_port_lists_dep = []
for d, de, dep in zip(last_edit_index_port, t_ct_at_last_edit, uport_list_by_seq):
    zip_list = list(zip(d, de, dep))
    zipped_port_lists_dep.append(zip_list)

# sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
uport_list_by_seq_sorted_pre_dep = []
for lis in zipped_port_lists_dep:
    sorted_list = sorted(lis, key=lambda x: (-x[0], x[1]))
    uport_list_by_seq_sorted_pre_dep.append(sorted_list)

# Extract the sorted items
uport_list_by_seq_sorted_dep = []
for dep in uport_list_by_seq_sorted_pre_dep:
    selection = [item[2] for item in dep]
    uport_list_by_seq_sorted_dep.append(selection)

uport_seqs_sorted2_dep = []
for dep in uport_list_by_seq_sorted_dep:
    uport_seqs_sorted2_dep_temp = []
    for seq_list in dep:
        uport_seqs_sorted2_dep_temp.append(seq_list[0][0])
    uport_seqs_sorted2_dep.append(uport_seqs_sorted2_dep_temp)

uport_seqs_sorted_dep = []
for dep in uport_seqs_sorted2_dep:
    uport_seqs_sorted_dep_temp = []
    for sequence in dep:
        uport_seqs_sorted_dep_temp.append(sequence[:-revprimlen2])
    uport_seqs_sorted_dep.append(uport_seqs_sorted_dep_temp)

# determine the order of sequences in the case that you sort by depth
dep_order_list_port = []
for dep, dep2 in zip(uport_sequences_og_order, uport_seqs_sorted_dep):
    dep_order_list_port_temp = []
    for seq2 in dep:
        for i, sequence in enumerate(dep2):
            if sequence == seq2:
                dep_order_list_port_temp.append(i + 1)
                break
    dep_order_list_port.append(dep_order_list_port_temp)

## processing data for portioned sequence difference ##
# extract the normalized frequencies of each sequence for your control sample
control_uport_seq_frequencies = []
for dep in uport_list_by_seq:
    control_uport_seq_frequencies_temp = []
    for cir in dep:
        control_uport_seq_frequencies_temp.append(int(cir[control_index][3]))
    control_uport_seq_frequencies.append(control_uport_seq_frequencies_temp)

# extract the frequencies of the non-controls
all_uport_seq_frequencies = []
for dep in uport_list_by_seq:
    all_uport_seq_frequencies_temp = []
    for cir in dep:
        non_seq_frequencies_temp = []
        for mut in cir:
            if isinstance(mut, list):
                non_seq_frequencies_temp.append(mut[3])
            else:
                non_seq_frequencies_temp.append(0)
        all_uport_seq_frequencies_temp.append(non_seq_frequencies_temp)
    all_uport_seq_frequencies.append(all_uport_seq_frequencies_temp)

# find the absolute difference between the number of reads in the control sample and the rest. Also determine total difference across mutants
total_seq_dif_port = []
seq_dif_port = []
for dep1, dep2 in zip(all_uport_seq_frequencies, control_uport_seq_frequencies):
    total_seq_dif_port_temp = []
    seq_dif_port_temp = []
    for cir_ex, cir_con in zip(dep1, dep2):
        mut_dif = []
        for mut in cir_ex:
            mut_dif.append(abs(mut - cir_con))
        total_seq_dif_port_temp.append(sum(mut_dif))
        seq_dif_port_temp.append(mut_dif)
    total_seq_dif_port.append(total_seq_dif_port_temp)
    seq_dif_port.append(seq_dif_port_temp)

combined_uport_by_dif = []
for d1, d2 in zip(total_seq_dif_port, uport_list_by_seq):
    zip_list = list(zip(d1, d2))
    combined_uport_by_dif.append(zip_list)

uport_sorted_by_dif_pre = []
for d1 in combined_uport_by_dif:
    sorted_list = sorted(d1, key=lambda x: x[0], reverse=True)
    uport_sorted_by_dif_pre.append(sorted_list)

uport_list_by_seq_sorted_dif = []
for dep in uport_sorted_by_dif_pre:
    sorted_list = [x[1] for x in dep]
    uport_list_by_seq_sorted_dif.append(sorted_list)

uport_seqs_sorted2_dif = []
for dep in uport_list_by_seq_sorted_dif:
    uport_seqs_sorted2_dif_temp = []
    for seq_list in dep:
        uport_seqs_sorted2_dif_temp.append(seq_list[0][0])
    uport_seqs_sorted2_dif.append(uport_seqs_sorted2_dif_temp)

uport_seqs_sorted_dif = []
for dep in uport_seqs_sorted2_dif:
    uport_seqs_sorted_dif_temp = []
    for sequence in dep:
        uport_seqs_sorted_dif_temp.append(sequence[:-revprimlen2])
    uport_seqs_sorted_dif.append(uport_seqs_sorted_dif_temp)

# determine the order of sequences in the case that you sort by difference
dif_order_list_port = []
for dep1, dep2 in zip(uport_sequences_og_order, uport_seqs_sorted_dif):
    dif_order_list_port_temp = []
    for seq2 in dep1:
        for i, sequence in enumerate(dep2):
            if sequence == seq2:
                dif_order_list_port_temp.append(i + 1)
                break
    dif_order_list_port.append(dif_order_list_port_temp)

combined_uport_by_dif_final = []
for a, b, c, d, e in zip(total_seq_dif_port, uport_list_by_seq, dep_order_list_port, dif_order_list_port, seq_dif_port):
    zist = list(zip(a, b, c, d, e))
    combined_uport_by_dif_final.append(zist)
zipped_port_lists_dep_final = []
for a, b, c, d, e, f in zip(last_edit_index_port, t_ct_at_last_edit, uport_list_by_seq, dep_order_list_port, dif_order_list_port, seq_dif_port):
    zist = list(zip(a, b, c, d, e, f))
    zipped_port_lists_dep_final.append(zist)

if sort_choice == "depth":
    print("Sorting sequence portions by editing depth.")

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    uport_list_by_seq_sorted_pre = []
    for dep in zipped_port_lists_dep_final:
        sist = sorted(dep, key=lambda x: (-x[0], x[1]))
        uport_list_by_seq_sorted_pre.append(sist)

    # Extract the sorted items
    uport_list_by_seq_sorted = []
    uport_depth_order_list_sorted = []
    uport_dif_order_list_sorted = []
    uport_dif_per_mut_sorted = []
    for dep in uport_list_by_seq_sorted_pre:
        uport_list_by_seq_sorted.append([item[2] for item in dep])
        uport_depth_order_list_sorted.append([item[3] for item in dep])
        uport_dif_order_list_sorted.append([item[4] for item in dep])
        uport_dif_per_mut_sorted.append([item[5] for item in dep])

    uport_seqs_sorted2 = []
    for dep in uport_list_by_seq_sorted:
        uport_seqs_sorted2_temp = []
        for seq_list in dep:
            uport_seqs_sorted2_temp.append(seq_list[0][0])
        uport_seqs_sorted2.append(uport_seqs_sorted2_temp)

    uport_seqs_sorted = []
    for dep in uport_seqs_sorted2:
        uport_seqs_sorted_temp = []
        for sequence in dep:
            uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
        uport_seqs_sorted.append(uport_seqs_sorted_temp)

elif sort_choice == "difference":
    print("Sorting sequence portions by difference between mutant and control.")

    uport_sorted_by_dif_pre = []
    for dep in combined_uport_by_dif_final:
        sist = sorted(dep, key=lambda x: x[0], reverse=True)
        uport_sorted_by_dif_pre.append(sist)

    uport_list_by_seq_sorted = []
    uport_depth_order_list_sorted = []
    uport_dif_order_list_sorted = []
    uport_dif_per_mut_sorted = []
    for dep in uport_sorted_by_dif_pre:
        uport_list_by_seq_sorted.append([x[1] for x in dep])
        uport_depth_order_list_sorted.append([item[2] for item in dep])
        uport_dif_order_list_sorted.append([item[3] for item in dep])
        uport_dif_per_mut_sorted.append([item[4] for item in dep])

    uport_seqs_sorted2 = []
    for dep in uport_list_by_seq_sorted:
        uport_seqs_sorted2_temp = []
        for seq_list in dep:
            uport_seqs_sorted2_temp.append(seq_list[0][0])
        uport_seqs_sorted2.append(uport_seqs_sorted2_temp)

    uport_seqs_sorted = []
    for dep in uport_seqs_sorted2:
        uport_seqs_sorted_temp = []
        for sequence in dep:
            uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
        uport_seqs_sorted.append(uport_seqs_sorted_temp)

else:
    print("There may be a typo in your choice of sequence sorting option. The script will continue without intentional sorting of portioned sequences unless terminated.")

    uport_list_by_seq_sorted = uport_list_by_seq

    uport_seqs_sorted2 = []
    for dep in uport_list_by_seq_sorted:
        uport_seqs_sorted2_temp = []
        for seq_list in dep:
            uport_seqs_sorted2_temp.append(seq_list[0][0])
        uport_seqs_sorted2.append(uport_seqs_sorted2_temp)

    uport_seqs_sorted = []
    for dep in uport_seqs_sorted2:
        uport_seqs_sorted_temp = []
        for sequence in dep:
            uport_seqs_sorted_temp.append(sequence[:-revprimlen2])
        uport_seqs_sorted.append(uport_seqs_sorted_temp)

    uport_depth_order_list_sorted = dep_order_list_port
    uport_dif_order_list_sorted = dif_order_list_port
    uport_dif_per_mut_sorted = seq_dif_port

##### end of the alignment for the sequences subset by editing site depth #####

# take out sequences for final pe and fe alignments
finally_pe2 = finally_aligned_seqs2[0]
finally_fe2 = finally_aligned_seqs2[1]
finally_pe22 = finally_aligned_seqs22[0]
finally_fe22 = finally_aligned_seqs22[1]

# create a list containing unique sequences, excluding pe and fe
unique_sequences = finally_aligned_seqs2[2:]

# add the final aligned seqs to each adapted_list sublist
for sublist in adapted_list:
    for subber, sequences in zip(sublist, unique_sequences):
        if isinstance(subber, list):
            subber.append(sequences)

adapted_list_by_seq = list(map(list, zip(*adapted_list)))

# take the processed sequences in their original order
adapted_seqs = []
for seq_list in adapted_list_by_seq:
    for mut in seq_list:
        if isinstance(mut, list):
            adapted_seqs.append(mut[4])
            break

# add normalized read counts to each mutant sublist
for sequence in adapted_list_by_seq:
    for mut, rpmr in zip(sequence, rpmr_scaling):
        if isinstance(mut, list):
            normalized_ct = int(round(int(mut[2]) * rpmr, 0))
            mut.append(normalized_ct)


# take out the t counts in the pe sequence for use as a reference
pe_t_ct = t_counts_new2[0]
# create a list containing the last edited site by comparing T counts to PE reference. If no editing occurs, len of amplicon is used
last_edit_index = []
for sequence in t_counts_new2[2:]:
    id_list_temp = []
    for i, (loc, ref) in enumerate(zip(sequence, pe_t_ct)):
        if loc != ref:
            id_list_temp.append(i)
    if len(id_list_temp) > 0:
        last_edit_index.append(id_list_temp[0])
    else:
        last_edit_index.append(len(sequence) - 1)

# create an arbitrary ranking system
t_ct_at_last_edit_all_uq = []
for leip, tcpu in zip(last_edit_index, t_counts_new2[2:]):
    t_ct_at_last_edit_all_uq.append(int(tcpu[leip]))

# zip the lists together
zipped_lists_all = list(zip(last_edit_index, t_ct_at_last_edit_all_uq, adapted_list_by_seq))

# sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
list_by_seq_sorted_pre = sorted(zipped_lists_all, key=lambda x: (-x[0], x[1]))

# Extract the sorted items
list_by_seq_sorted_dep = [item[2] for item in list_by_seq_sorted_pre]
unique_seqs_sorted_dep = []
for seq_list in list_by_seq_sorted_dep:
    for mut in seq_list:
        if isinstance(mut, list):
            unique_seqs_sorted_dep.append(mut[4])
            break

# find the final order of the sequences when sorted by depth while maintaining input sequence order
depth_order_list = []
for seq2 in adapted_seqs:
    for i, sequence in enumerate(unique_seqs_sorted_dep):
        if sequence == seq2:
            depth_order_list.append(i + 1)
            break

# extract the normalized frequencies of each sequence for your control sample
control_seq_frequencies = []
for cir in adapted_list_by_seq:
    if isinstance(cir[control_index], list):
        control_seq_frequencies.append(int(cir[control_index][6]))
    else:
        control_seq_frequencies.append(0)

# extract the frequencies of the non-controls
all_seq_frequencies = []
for cir in adapted_list_by_seq:
    non_seq_frequencies_temp = []
    for mut in cir:
        if isinstance(mut, list):
            non_seq_frequencies_temp.append(mut[6])
        else:
            non_seq_frequencies_temp.append(0)
    all_seq_frequencies.append(non_seq_frequencies_temp)

# find the absolute difference between the number of reads in the control sample and the rest. Also determine total difference across mutants
total_seq_dif = []
dif_per_mut = []
for cir_ex, cir_con in zip(all_seq_frequencies, control_seq_frequencies):
    mut_dif = []
    for mut in cir_ex:
        mut_dif.append(abs(mut - cir_con))
    total_seq_dif.append(sum(mut_dif))
    dif_per_mut.append(mut_dif)

combined_by_dif = list(zip(total_seq_dif, adapted_list_by_seq))

list_sorted_by_dif_pre = sorted(combined_by_dif, key=lambda x: x[0], reverse=True)

list_by_seq_sorted_dif = [x[1] for x in list_sorted_by_dif_pre]

unique_seqs_sorted_dif = []
for seq_list in list_by_seq_sorted_dif:
    for mut in seq_list:
        if isinstance(mut, list):
            unique_seqs_sorted_dif.append(mut[4])
            break

# determine the order of sequences in the case that you sort by difference
dif_order_list = []
for seq2 in adapted_seqs:
    for i, sequence in enumerate(unique_seqs_sorted_dif):
        if sequence == seq2:
            dif_order_list.append(i + 1)
            break

#####  start of sorting by editing depth #####
if sort_choice == "depth":
    print("Sorting whole sequences by editing depth.")

    # zip the lists together and sort everything by depth
    zipped_lists_all2 = list(zip(last_edit_index, t_ct_at_last_edit_all_uq, adapted_list_by_seq, depth_order_list, dif_order_list, dif_per_mut))

    # sort based on the first list (priority_list) in descending order, in case of ties, use amount of Us at that site
    list_by_seq_sorted_pre = sorted(zipped_lists_all2, key=lambda x: (-x[0], x[1]))

    # Extract the sorted items
    list_by_seq_sorted = [item[2] for item in list_by_seq_sorted_pre]
    depth_order_list_sorted = [item[3] for item in list_by_seq_sorted_pre]
    dif_order_list_sorted = [item[4] for item in list_by_seq_sorted_pre]
    dif_per_mut_sorted = [item[5] for item in list_by_seq_sorted_pre]

    unique_seqs_sorted = []
    for seq_list in list_by_seq_sorted:
        for mut in seq_list:
            if isinstance(mut, list):
                unique_seqs_sorted.append(mut[5])
                break

##### end of sorting by editing depth, start of sorting by difference between mutant and control #####
elif sort_choice == "difference":
    print("Sorting whole sequences by difference between mutant and control.")

    # zip the lists together and sort everything by difference
    combined_by_dif2 = list(zip(total_seq_dif, adapted_list_by_seq, depth_order_list, dif_order_list))

    # sort the list based on total difference between sequences
    list_sorted_pre = sorted(combined_by_dif2, key=lambda x: x[0], reverse=True)

    list_by_seq_sorted = [x[1] for x in list_sorted_pre]
    depth_order_list_sorted = [x[2] for x in list_sorted_pre]
    dif_order_list_sorted = [x[3] for x in list_sorted_pre]
    dif_per_mut_sorted = [item[4] for item in list_sorted_pre]

    unique_seqs_sorted = []
    for seq_list in list_by_seq_sorted:
        for mut in seq_list:
            if isinstance(mut, list):
                unique_seqs_sorted.append(mut[5])
                break

#### this commented chunk will re-arrange the list so the control is the first column, but we will need to rearrange other lists e.g. mutant names if we want to use this #####
    # control_list = []
    # for sequ in list_by_seq_sorted:
    #     control_list.append(sequ[control_index])
    #
    # experimental_samp_list = []
    # for sequ in list_by_seq_sorted:
    #     temp_list = []
    #     for i, seu in enumerate(sequ):
    #         if i == control_index:
    #             continue
    #         else:
    #             temp_list.append(seu)
    #     experimental_samp_list.append(temp_list)
    #
    # list_by_seq_sorted = []
    # for con, exp_list in zip(control_list, experimental_samp_list):
    #     new_list = [con] + exp_list
    #     list_by_seq_sorted.append(new_list)

##### end of sorting by difference between mutant and control #####
else:
    print("There is a typo in your choice of sequence sorting option. The script will continue without intentional sorting of whole sequences unless terminated.")

    list_by_seq_sorted = adapted_list_by_seq

    unique_seqs_sorted = []
    for seq_list in list_by_seq_sorted:
        unique_seqs_sorted.append(seq_list[0][5])

    depth_order_list_sorted = depth_order_list
    dif_order_list_sorted = dif_order_list
    dif_per_mut_sorted = dif_per_mut

# determine the lengths of each sequence
seq_lengths = []
for seek in unique_seqs_sorted:
    no_dash = seek.replace("-", "")
    no_ast = no_dash.replace("*", "")
    leng = len(no_ast)
    seq_lengths.append(leng)

# process the sequences so non-U stretches are pasted in as periods
print("Simplifying sequences by subbing non-U nucleotides for periods...")
period_pattern = re.compile(r'[ACG]+')
period_seqs = []
for seeq in unique_seqs_sorted:
    output_string = period_pattern.sub('.', seeq)
    period_seqs.append(output_string)
pe_period = period_pattern.sub(".", finally_pe2)
fe_period = period_pattern.sub(".", finally_fe2)

# extract the total percentages of sequences pulled for each cell line
perc_sums = []
for i, val in enumerate(mutant_names):
    mut_perc_sum = []
    for sublist in list_by_seq_sorted:
        if isinstance(sublist[i], str):
            continue
        else:
            final = round(float(sublist[i][4]) / 100, 4)
            mut_perc_sum.append(final)
    sumn = sum(mut_perc_sum)
    perc_sums.append(sumn)

# repeat the above for absolute number of reads
abs_count_sums = []
for i, val in enumerate(mutant_names):
    mut_perc_sum = []
    for sublist in list_by_seq_sorted:
        if isinstance(sublist[i], str):
            continue
        else:
            final = int(sublist[i][2])
            mut_perc_sum.append(final)
    sumn = sum(mut_perc_sum)
    abs_count_sums.append(sumn)

# repeat the above for reads per million reads
rpmr_sums = []
for i, val in enumerate(mutant_names):
    mut_perc_sum = []
    for sublist in list_by_seq_sorted:
        if isinstance(sublist[i], str):
            continue
        else:
            final = int(sublist[i][6])
            mut_perc_sum.append(final)
    sumn = sum(mut_perc_sum)
    rpmr_sums.append(sumn)

## repeat the above functions for the portioned sequences
port_perc_sums = []
for dep in uport_list_by_seq_sorted:
    mut_perc_sum = []
    for i, val in enumerate(mutant_names):
        mut_perc_sum_temp = []
        for sublist in dep:
            if isinstance(sublist[i], str):
                continue
            else:
                final = round(float(sublist[i][1]) / 100, 4)
                mut_perc_sum_temp.append(final)
        sumn = sum(mut_perc_sum_temp)
        mut_perc_sum.append(sumn)
    port_perc_sums.append(mut_perc_sum)

# repeat the above for absolute number of reads
port_abs_count_sums = []
for dep in uport_list_by_seq_sorted:
    mut_perc_sum = []
    for i, val in enumerate(mutant_names):
        mut_perc_sum_temp = []
        for sublist in dep:
            if isinstance(sublist[i], str):
                continue
            else:
                final = int(sublist[i][2])
                mut_perc_sum_temp.append(final)
        sumn = sum(mut_perc_sum_temp)
        mut_perc_sum.append(sumn)
    port_abs_count_sums.append(mut_perc_sum)

# repeat the above for reads per million reads
port_rpmr_sums = []
for dep in uport_list_by_seq_sorted:
    mut_perc_sum = []
    for i, val in enumerate(mutant_names):
        mut_perc_sum_temp = []
        for sublist in dep:
            if isinstance(sublist[i], str):
                continue
            else:
                final = int(sublist[i][3])
                mut_perc_sum_temp.append(final)
        sumn = sum(mut_perc_sum_temp)
        mut_perc_sum.append(sumn)
    port_rpmr_sums.append(mut_perc_sum)


print("All necessary data has been processed. Beginning to write out alignments to Excel sheet.")

# lists for sorting headers and data
mutant_names1 = ["Sequence", "Match?", "Length", "Depth Order", "Difference Order", ""] + mutant_names
mutant_names2 = ["Sequence", "Match?", "Length", "Depth Order", "Difference Order", ""] + mutant_names + [""] + mutant_names + [""] + mutant_names
mutant_names_period = ["Sequence", "Match?", "Length", "Depth Order", "Difference Order", ""] + mutant_names + [""] + mutant_names + [""] + mutant_names

# create excel workbook
excel_output = inter_analysis_folder + "/all_cell_lines_alignments_" + gene + "_" + today + ".xlsx"
workbook = xlsxwriter.Workbook(excel_output)
ws1 = workbook.add_worksheet("Full Sequence")
ws3_name = "Period Seq Alignment"
ws3 = workbook.add_worksheet(ws3_name)
portion_length = sorted([i for i in range(1, depth2 + 1)], key=lambda x: x, reverse=True)
ws2_name = []
for i, (num, dep) in enumerate(zip(editing_depth, portion_length)):
    sheet_name = "Edit site " + num.upper() + " (depth = " + str(dep) + " ES)"
    ws2_name.append(sheet_name)
sheet_calls = [None] * len(ws2_name) # intialize list for the portioned tabs
sheet_calls_list = [] # store the sheet calls for the portions
for i, name in enumerate(ws2_name):
    sheet_calls[i] = workbook.add_worksheet(name)
    sheet_calls_list.append(sheet_calls[i])
ws5 = workbook.add_worksheet("Difference Matrix")
ws4 = workbook.add_worksheet("QC Info")
worksheets = [ws1, ws3] + sheet_calls_list
script_name = os.path.basename(sys.argv[0])
today = datetime.today().strftime('%m-%d-%Y')


# define some formats
green_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#33CC33', 'bold': True})
purple_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#3304FC', 'bold': True})
red_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#FF0000', 'bold': True})
cyan_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#3304FC', 'bold': True})
default_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#000000'})  # Default color (black)
bold_format = workbook.add_format({'font_name': 'Courier New', 'font_color': '#000000', 'bold': True})
header_format = workbook.add_format({'font_name': 'Arial', 'align': 'center'})
basic_format = workbook.add_format({'font_name': 'Arial'})
merge_format = workbook.add_format({
    "align": "center",
    "valign": "vcenter",
    'font_name': 'Arial'
})

# prepare some columnn headers for each tab containing sequences
for ws in worksheets:
    ws.merge_range(0, 6, 0, len(mutant_names) + 5, "Percent Frequencies", merge_format)
    ws.merge_range(0, len(mutant_names1) + 1, 0, len(mutant_names1) + len(mutant_names), "Absolute # of Reads", merge_format)
    ws.merge_range(0, len(mutant_names1) + 2 + len(mutant_names), 0, len(mutant_names2) - 1, "Reads per Million Reads", merge_format)

# write out whole sequence data to the first sheet in the workbook
for col_index, names in enumerate(mutant_names2):
    ws1.write(1, col_index, names, header_format)

for row_index, (custom_alignment, lol, lengs, deppy, diffy) in enumerate(zip(unique_seqs_sorted, list_by_seq_sorted, seq_lengths, depth_order_list_sorted, dif_order_list_sorted)):

    # Prepare the rich string with formatting
    rich_string = []

    if custom_alignment == finally_pe2:
        ws1.write(row_index + 2, 1, "Pre-edited", header_format)
        ws1.write(row_index + 2, 0, custom_alignment, default_format)
    else:
        # write out the sequences with color coding to the appropriate cell
        for character, pong, fong in zip(custom_alignment, finally_pe2, finally_fe2):
            if character == '*' and pong == "U" and fong == "U":
                rich_string.append(green_format)
            elif character == "u" and fong == "-":
                rich_string.append(purple_format)
            elif character == "u" and pong == "-" and fong == "u":
                rich_string.append(red_format)
            elif character == "U" and pong == "U" and fong == "*":
                rich_string.append(cyan_format)
            elif character == "*" and pong == "U" and fong == "*":
                rich_string.append(bold_format)
            else:
                rich_string.append(default_format)
            rich_string.append(character)
        # Write the rich string to the cell
        ws1.write_rich_string(row_index + 2, 0, *rich_string)

    mut_count_abs = []
    mut_perc = []
    mut_count_norm = []
    for seq_item, norms in zip(lol, rpmr_scaling):
        if isinstance(seq_item, list):
            mut_count_abs.append(int(seq_item[2]))
            mut_perc.append(round(float(seq_item[4]), 2) / 100)
            mut_count_norm.append(round(int(seq_item[2]) * norms, 0))
        else:
            mut_count_abs.append("n.d.")
            mut_perc.append("n.d.")
            mut_count_norm.append("n.d.")

    ws1.write(row_index + 2, 2, lengs, header_format)
    ws1.write(row_index + 2, 3, deppy, header_format)
    ws1.write(row_index + 2, 4, diffy, header_format)

    for col_index2, mp in enumerate(mut_perc):
        ws1.write(row_index + 2, col_index2 + 6, mp, header_format)
    for col_index3, ps in enumerate(perc_sums):
        ws1.write(len(list_by_seq_sorted) + 2, col_index3 + 6, ps, header_format)

    for col_index2, mc in enumerate(mut_count_abs):
        ws1.write(row_index + 2, col_index2 + len(mutant_names2) - (len(mutant_names) * 2) - 1, mc, header_format)
    for col_index3, acs in enumerate(abs_count_sums):
        ws1.write(len(list_by_seq_sorted) + 2, col_index3 + len(mutant_names2) - (len(mutant_names) * 2) - 1, int(acs), header_format)

    for col_index2, mcn in enumerate(mut_count_norm):
        ws1.write(row_index + 2, col_index2 + len(mutant_names2) - len(mutant_names), mcn, header_format)
    for col_index3, rs in enumerate(rpmr_sums):
        ws1.write(len(list_by_seq_sorted) + 2, col_index3 + len(mutant_names2) - len(mutant_names), int(rs), header_format)

    if custom_alignment == finally_fe2:
        ws1.write(row_index + 2, 1, "Fully edited", header_format)

# write out the portioned data
for (wsw, one, two, three, four, five, six, seven, eight) in zip(sheet_calls_list, uport_seqs_sorted, uport_list_by_seq_sorted, uport_depth_order_list_sorted, uport_dif_order_list_sorted,
                                                          port_perc_sums, port_abs_count_sums, port_rpmr_sums, proc_index):
    for col_index, names in enumerate(mutant_names2):
        wsw.write(1, col_index, names, header_format)

    for row_index, (custom_alignment, lol, deppy, diffy) in enumerate(zip(one, two, three, four)):

        cu_no_dash = custom_alignment.replace("-", "")
        cu_no_ast = cu_no_dash.replace("*", "")
        leg = len(cu_no_ast)
        wsw.write(row_index + 2, 2, leg, header_format)
        wsw.write(row_index + 2, 3, deppy, header_format)
        wsw.write(row_index + 2, 4, diffy, header_format)

        # Prepare the rich string with formatting
        rich_string = []

        if custom_alignment == finally_pe2[eight:-revprimlen2]:
            wsw.write(row_index + 2, 1, "Matches pre-edited start", header_format)
            wsw.write(row_index + 2, 0, custom_alignment, default_format)
        else:
            # write out the sequences with color coding to the appropriate cell
            for character, pong, fong in zip(custom_alignment, finally_pe22[eight:-revprimlen2], finally_fe22[eight:-revprimlen2]):
                if character == '*' and pong == "U" and fong == "U":
                    rich_string.append(green_format)
                elif character == "u" and fong == "-":
                    rich_string.append(purple_format)
                elif character == "u" and pong == "-" and fong == "u":
                    rich_string.append(red_format)
                elif character == "U" and pong == "U" and fong == "*":
                    rich_string.append(cyan_format)
                elif character == "*" and pong == "U" and fong == "*":
                    rich_string.append(bold_format)
                else:
                    rich_string.append(default_format)
                rich_string.append(character)
            # Write the rich string to the cell
            wsw.write_rich_string(row_index + 2, 0, *rich_string)

        mut_count = []
        mut_perc = []
        mut_count_norm_port = []
        for seq_item, norms in zip(lol, rpmr_scaling):
            if seq_item[2] > 0:
                mut_perc.append(seq_item[1] / 100)
                mut_count.append(int(seq_item[2]))
                mut_count_norm_port.append(int(round(seq_item[2] * norms, 0)))
            else:
                mut_perc.append("n.d.")
                mut_count.append("n.d.")
                mut_count_norm_port.append("n.d.")

        for col_index2, mp in enumerate(mut_perc):
            wsw.write(row_index + 2, col_index2 + 6, mp, header_format)
        for col_index3, ps in enumerate(five):
            wsw.write(len(one) + 2, col_index3 + 6, ps, header_format)

        for col_index2, mc in enumerate(mut_count):
            wsw.write(row_index + 2, col_index2 + len(mutant_names2) - (len(mutant_names) * 2) - 1, mc, header_format)
        for col_index3, acs in enumerate(six):
            wsw.write(len(one) + 2, col_index3 + len(mutant_names2) - (len(mutant_names) * 2) - 1, int(acs), header_format)

        for col_index2, mcn in enumerate(mut_count_norm_port):
            wsw.write(row_index + 2, col_index2 + len(mutant_names2) - len(mutant_names), mcn, header_format)
        for col_index3, rs in enumerate(seven):
            wsw.write(len(one) + 2, col_index3 + len(mutant_names2) - len(mutant_names), int(rs), header_format)

        if custom_alignment == finally_fe2[eight:-revprimlen2]:
            wsw.write(row_index + 2, 1, "Matches fully edited start", header_format)

# write out period sequence data to the third sheet in the workbook
for col_index, names in enumerate(mutant_names_period):
    ws3.write(1, col_index, names, header_format)

for row_index, (custom_alignment, lol, deppy, diffy) in enumerate(zip(period_seqs, list_by_seq_sorted, depth_order_list_sorted, dif_order_list_sorted)):

    up_u_count = custom_alignment.count("U")
    low_u_count = custom_alignment.count("u")
    u_sum = up_u_count + low_u_count
    ws3.write(row_index + 2, 2, u_sum, header_format)
    ws3.write(row_index + 2, 3, deppy, header_format)
    ws3.write(row_index + 2, 4, diffy, header_format)

    # Prepare the rich string with formatting
    rich_string = []

    if custom_alignment == pe_period:
        ws3.write(row_index + 2, 1, "Pre-edited", header_format)
        ws3.write(row_index + 2, 0, custom_alignment, default_format)
    else:
        # write out the sequences with color coding to the appropriate cell
        for character, pong, fong in zip(custom_alignment, pe_period, fe_period):
            if character == '*' and pong == "U" and fong == "U":
                rich_string.append(green_format)
            elif character == "u" and fong == "-":
                rich_string.append(purple_format)
            elif character == "u" and pong == "-" and fong == "u":
                rich_string.append(red_format)
            elif character == "U" and pong == "U" and fong == "*":
                rich_string.append(cyan_format)
            elif character == "*" and pong == "U" and fong == "*":
                rich_string.append(bold_format)
            else:
                rich_string.append(default_format)
            rich_string.append(character)
        # Write the rich string to the cell
        ws3.write_rich_string(row_index + 2, 0, *rich_string)

    mut_count = []
    mut_perc = []
    mut_count_norm = []
    for seq_item, norms in zip(lol, rpmr_scaling):
        if isinstance(seq_item, list):
            mut_count.append(int(seq_item[2]))
            mut_perc.append(round(float(seq_item[4]), 2) / 100)
            mut_count_norm.append(round(int(seq_item[2]) * norms, 0))
        else:
            mut_count.append("n.d.")
            mut_perc.append("n.d.")
            mut_count_norm.append("n.d.")

    for col_index2, mp in enumerate(mut_perc):
        ws3.write(row_index + 2, col_index2 + 6, mp, header_format)
    for col_index3, ps in enumerate(perc_sums):
        ws3.write(len(list_by_seq_sorted) + 2, col_index3 + 6, ps, header_format)

    for col_index2, mc in enumerate(mut_count):
        ws3.write(row_index + 2, col_index2 + len(mutant_names_period) - (len(mutant_names) * 2) - 1, mc, header_format)
    for col_index3, acs in enumerate(abs_count_sums):
        ws3.write(len(list_by_seq_sorted) + 2, col_index3 + len(mutant_names2) - (len(mutant_names) * 2) - 1, int(acs), header_format)

    for col_index2, mcn in enumerate(mut_count_norm):
        ws3.write(row_index + 2, col_index2 + len(mutant_names_period) - len(mutant_names), mcn, header_format)
    for col_index3, rs in enumerate(rpmr_sums):
        ws3.write(len(list_by_seq_sorted) + 2, col_index3 + len(mutant_names2) - len(mutant_names), int(rs), header_format)

    if custom_alignment == fe_period:
        ws3.write(row_index + 2, 1, "Fully edited", header_format)

# fill out the QC tab
col_headers = ["# Top sequences pulled", "# of cell lines", "Control Sample", "Sorting method", "Date ran", "Script version"]
qc_data = [cir_depth, len(cir_files), mutant_names[control_index], sort_choice, today, script_name]
for row_index, title in enumerate(col_headers):
    ws4.write(row_index, 0, title, header_format)

for row_index, info in enumerate(qc_data):
    ws4.write(row_index, 1, info, header_format)

# add the difference matrices to a new sheet
ws5.merge_range(0, 0, 0, len(mutant_names) - 1, "Full Sequence", merge_format)
merge_stat_p = len(mutant_names) + 1
merge_stat_l = len(mutant_names) - 1
for i, (dep, dif_dep) in enumerate(zip(editing_depth, uport_dif_per_mut_sorted)):
    start = merge_stat_p * (i + 1)
    end = merge_stat_p * (i + 1) + merge_stat_l
    ws5.merge_range(0, start, 0, end, "Portioned Sequence: " + dep.upper(), merge_format)
    for col_index, names in enumerate(mutant_names):
        ws5.write(1, start + col_index, names, header_format)
    for row_index, seq in enumerate(dif_dep):
        for col_index2, mut in enumerate(seq):
            ws5.write(row_index + 2, start + col_index2, mut, header_format)

for col_index, names in enumerate(mutant_names):
    ws5.write(1, col_index, names, header_format)

for row_index, seq in enumerate(dif_per_mut_sorted):
    for col_index, mut in enumerate(seq):
        ws5.write(row_index + 2, col_index, mut, header_format)


# save the workbook
workbook.close()


# generate sequence diversity data
seq_diversity_list = []
for mut, trl, nir in zip(mutant_names, total_reads_list, num_identical_reads):
    unique_perc = round((nir / trl) * 100, 2)
    new_list = [mut, trl, nir, unique_perc]
    seq_diversity_list.append(new_list)

# create sequence diversity excel book
seq_diversity_excel = inter_analysis_folder + "/" + gene + "_sequence_diversity_" + today + ".xlsx"
workbook = xlsxwriter.Workbook(seq_diversity_excel)
ws10 = workbook.add_worksheet()
seq_div_col_names = ["Cell Line", "Total Reads", "# Unique Reads", "% Unique Reads"]

# write the headers out to the file
for col_index, head in enumerate(seq_div_col_names):
    ws10.write(0, col_index, head, bold_format)

# write the data out to the file
for row_index, mut in enumerate(seq_diversity_list):
    for col_index, dat in enumerate(mut):
        ws10.write(row_index + 1, col_index, dat, header_format)

workbook.close()
